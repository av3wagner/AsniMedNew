from IPython.core.display import HTML 
from collections import Counter 
from colorama import Fore, Style # maakes strings colored 
from matplotlib import * 
from matplotlib import pyplot as plt 
from matplotlib.colors import ListedColormap 
from plotly import tools 
from plotly.offline import download_plotlyjs, init_notebook_mode, plot, iplot 
from plotly.offline import iplot 
from plotly.subplots import make_subplots 
from sklearn import datasets, linear_model, metrics 
from sklearn import preprocessing 
from sklearn.compose import make_column_transformer 
from sklearn.discriminant_analysis import LinearDiscriminantAnalysis 
from sklearn.ensemble import RandomForestClassifier 
from sklearn.ensemble import RandomForestRegressor, RandomForestClassifier, GradientBoostingRegressor, ExtraTreesRegressor, AdaBoostClassifier
from sklearn.feature_selection import SelectKBest, SelectPercentile, f_classif, f_regression, mutual_info_regression
from sklearn.impute import SimpleImputer, KNNImputer 
from sklearn.linear_model import LogisticRegression 
from sklearn.metrics import accuracy_score, confusion_matrix, classification_report
from sklearn.model_selection import StratifiedKFold 
from sklearn.model_selection import cross_val_score 
from sklearn.model_selection import train_test_split 
from sklearn.naive_bayes import GaussianNB 
from sklearn.neighbors import KNeighborsClassifier 
from sklearn.neighbors import KNeighborsRegressor 
from sklearn.pipeline import Pipeline 
from sklearn.preprocessing import StandardScaler 
from sklearn.svm import SVC 
from sklearn.svm import SVR 
from sklearn.tree import DecisionTreeClassifier 
from sklearn.tree import plot_tree 
from termcolor import colored 
from xgboost import XGBRegressor, XGBClassifier 
from xgboost import plot_importance 
import colorama 
import cufflinks as cf 
import dash 
import matplotlib 
import matplotlib.pyplot as plt 
import numpy as np 
import pandas as pd 
import patchworklib as pw 
import plotly as pl 
import plotly as pplt 
import plotly.express as px 
import plotly.figure_factory as ff 
import plotly.graph_objects as go 
import plotly.io as pio 
import plotly.offline 
import plotly.offline as po 
import random 
import scipy.stats as stats 
import seaborn as sns 
import sys 
import warnings 
import time
#######################################
import os, sys, inspect, time, datetime
import subprocess
import json
from time import time, strftime, localtime
from datetime import timedelta
import shutil
from docx import Document
from configparser import ConfigParser
import configparser
from time import time, strftime, localtime
from datetime import timedelta
import pandas as pd
import pathlib

print("Start AsNiDefFa2.py!")

def AddBookmark(t, mark):    
    # Find a specific text or phrase in the document
    paragraph = section.AddParagraph()  
    paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
    text = paragraph.AppendText(t)
    text = document.FindString(t, False, True)
    # Get the found text as a single text range
    textRange = text.GetAsOneRange()
    # Get the paragraph where the text range is located
    paragraph = textRange.OwnerParagraph

    # Get the index position of the text in the paragraph
    index = paragraph.ChildObjects.IndexOf(textRange)

    # Add a bookmark start mark to the paragraph
    start = paragraph.AppendBookmarkStart(mark)
    # Insert the bookmark start mark at the index position of the text range
    paragraph.ChildObjects.Insert(index, start)
    # Add a bookmark end mark to the paragraph
    end = paragraph.AppendBookmarkEnd(mark)
    # Insert the bookmark end mark after the text range
    paragraph.ChildObjects.Insert(index + 2, end)

def boxplots_custom(dataset, columns_list, rows, cols, suptitle):
    fig, axs = plt.subplots(rows, cols, sharey=True, figsize=(13,5))
    fig.suptitle(suptitle,y=1, size=25)
    axs = axs.flatten()
    for i, data in enumerate(columns_list):
        sns.boxplot(data=dataset[data], orient='h', ax=axs[i])
        axs[i].set_title(data + ', skewness is: '+str(round(dataset[data].skew(axis = 0, skipna = True),2)))

def replace_zero_cholesterol(df):
    # Step 1: Create age groups and calculate average cholesterol for each group
    age_bins = [10, 20, 30, 40, 50, 60, 70, 80]
    age_labels = [f'{start}-{end}' for start, end in zip(age_bins[:-1], age_bins[1:])]
    df['AgeGroup'] = pd.cut(df['Age'], bins=age_bins, labels=age_labels, right=False)
    average_cholesterol_by_age = df.groupby('AgeGroup')['Cholesterol'].mean()

    # Step 2: Replace zero values in 'Cholesterol' with average values based on age groups
    def replace_zero(row):
        if row['Cholesterol'] == 0:
            return average_cholesterol_by_age[row['AgeGroup']]
        else:
            return row['Cholesterol']

    df['Cholesterol'] = df.apply(replace_zero, axis=1)

    # Drop the temporary 'AgeGroup' column
    df.drop(columns=['AgeGroup'], inplace=True)

#################### Def Section #######################
def getConfigFile(config):
    with open(config, encoding='utf-8') as json_file:
        return json.load(json_file)

def get_para_data(output_doc_name, paragraph):
    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
     
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
    output_para.paragraph_format.first_line_indent = Inches(0.25)
    output_para.paragraph_format.space_before = Pt(0)
    output_para.paragraph_format.space_after  = Pt(0)
    return output_para
    
########################################
def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def change_table_cell(cell, background_color=None, font_color=None, font_size=None, bold=None, italic=None):
    if background_color:
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    if font_color:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = docx.shared.RGBColor.from_string(font_color)

    if font_size:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = docx.shared.Pt(font_size)

    if bold is not None:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = bold

    if italic is not None:
        for p in cell.paragraphs:
            for r in p.runs:
                r.italic = italic


def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in ["top", "start", "bottom", "end"]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)
    
def add_table_to_doc_new(doc, df, heading, table_style='Table Grid', txt="EDA1"):
    for p in doc.paragraphs:
        if txt in p.text:
            doc.add_heading(heading, level=1).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ################## Start Table ####################
            columns = list(df.columns)
            print(columns)
            for col in columns: 
                df[col].fillna(" ", inplace = True) 
            
            # add table
            endcol=len(columns)
            table = doc.add_table(rows=1, cols=endcol, style=table_style)
            set_repeat_table_header(table.rows[0])
            table.autofit = True
            # add columns (if there is '_' then replace with space)
            for col in range(len(columns)):
                set_cell_margins(table.cell(0, col), top=100, start=100, bottom=100, end=50)
                table.cell(0, col).text = columns[col].replace("_", " ").capitalize()
            # add data
            for i, row in enumerate(df.itertuples()):
                table_row = table.add_row().cells
                for col in range(len(columns)):
                    #set_cell_margins(table_row[col], top=100, start=100, bottom=100, end=50)
                    set_cell_margins(table_row[col], top=10, start=10, bottom=10, end=5)
                    table_row[col].text = str(row[col+1])
                    if i==1:
                        col1=table.cell(i, 1).text 
                        nr=1
                        col2=table.cell(i, 2).text 
                    if table.cell(i, col).text=="nan":
                        table.cell(i, col).text==" "
                    #if i > 1 and col1==table.cell(i-1, col).text:
                        #table.cell(i, endcol).text=str(nr)
                        
            for row in range(df.shape[0]):
                for col in range(df.shape[-1]):
                    #table.cell(row+1, col).text = str(input_df.values[row, col])
                    #table.cell(row, col).width = widths[col]
                    #table.cell(row, col).hight = Cm(1)
                    table.cell(row+1, col).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT 
                    change_table_cell(table.rows[row+1].cells[col], background_color="lightgreen", font_color="0000ff", font_size=8, bold=True, italic=True)
                    #change_table_cell(table.rows[row+1].cells[col], background_color="White", font_color="000000", font_size=10, bold=False, italic=True)
                table.rows[row].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY  

        doc.add_section(WD_SECTION.ODD_PAGE)  
        return doc 

def move_table_after(table, paragraph):
    tbl = table._tbl 
    paragraph.add_run().element.addnext(tbl)

# Copy Table from Excel
def add_table_to_doc(doc, df, heading, table_style='Table Grid'):
    doc.add_heading(heading, level=1).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    columns = list(df.columns)
    print(columns)
    for col in columns: 
        df[col].fillna(" ", inplace = True) 
    
    # add table
    endcol=len(columns)
    table = doc.add_table(rows=1, cols=endcol, style=table_style)
    #set_repeat_table_header(table.rows[0])
    table.autofit = True
    # add columns (if there is '_' then replace with space)
    for col in range(len(columns)):
        set_cell_margins(table.cell(0, col), top=100, start=100, bottom=100, end=50)
        table.cell(0, col).text = columns[col].replace("_", " ").capitalize()
    # add data
    for i, row in enumerate(df.itertuples()):
        table_row = table.add_row().cells
        for col in range(len(columns)):
            #set_cell_margins(table_row[col], top=100, start=100, bottom=100, end=50)
            set_cell_margins(table_row[col], top=10, start=10, bottom=10, end=5)
            table_row[col].text = str(row[col+1])
            if i==1:
                col1=table.cell(i, 1).text 
                nr=1
                col2=table.cell(i, 2).text 
            if table.cell(i, col).text=="nan":
                table.cell(i, col).text==" "
            #if i > 1 and col1==table.cell(i-1, col).text:
                #table.cell(i, endcol).text=str(nr)
                
    for row in range(df.shape[0]):
        for col in range(df.shape[-1]):
            #table.cell(row+1, col).text = str(input_df.values[row, col])
            #table.cell(row, col).width = widths[col]
            #table.cell(row, col).hight = Cm(1)
            table.cell(row+1, col).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT 
            #change_table_cell(table.rows[row+1].cells[col], background_color="lightgreen", font_color="0000ff", font_size=8, bold=True, italic=True)
            change_table_cell(table.rows[row+1].cells[col], background_color="White", font_color="000000", font_size=10, bold=False, italic=True)
        table.rows[row].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY  
    
    doc.add_section(WD_SECTION.ODD_PAGE)  
    return doc

def tab_copy(k): 
    #table = form.tables[k]
    table = doc.tables[k]
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    paragraph = doc.add_paragraph("Hier wird die Tabelle №" + str(k+1) + " kopiert! Alle Daten wurde Simuliert!")
    paragraph._p.addnext(tbl)    

def set_column_width(table, column, width_mm):
    table.allow_autofit = False
    for row in table.rows:
        row.cells[column].width = Mm(width_mm)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def print_scores(y, y_pred):
    ac, pr, rc, f1 = accuracy_score(y, y_pred)*100, precision_score(y, y_pred)*100, recall_score(y, y_pred)*100, f1_score(y, y_pred, average='weighted')*100
    print(f"Accuracy:{ac}")
    print(f"Precision:{pr}")
    print(f"Recall:{rc}")
    print(f"F1-score:{f1}")
    return {'ac': ac, 'pr':pr, 'rc':rc, 'f1':f1}

def fig2img(fig): 
	buf = io.BytesIO() 
	fig.savefig(buf) 
	buf.seek(0) 
	img = Image.open(buf) 
	return img
    
def replace_copy0(txt):
    finder = wordapp.Selection.Find
    finder.Text = txt 
    finder.Execute()
    wordapp.Selection.MoveStart
    #wordapp.Selection.Paste()

def replace_copy(txt):
    finder = wordapp.Selection.Find
    finder.Text = txt 
    finder.Execute()
    wordapp.Selection.MoveStart
    wordapp.Selection.Paste()
    if txt=="KapEDA":
        time.sleep(10.0)
        
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

######################### Def for Test4RR.py ############################
def Text_copy(file):
    wordapp = win32com.client.gencache.EnsureDispatch("Word.Application")
    wordapp.Visible = False #True 
    worddoc = wordapp.Documents.Open(file)
    worddoc.Select() 
    wordapp.Selection.Copy()
    worddoc.ActiveWindow.Close()
    #wordapp.Application.Quit(-1)   

def update_tocHalb(doc):
    toc_count = doc.TablesOfContents.Count
    print(toc_count)
    stringG='INHALTSVERZEICHNIS'
    stringK='Содержание'
    if toc_count == 0:
        for i, p in enumerate(doc.Paragraphs):
            if stringK in p.Range.Text:
                try:
                    p.Range.InsertParagraphAfter()
                    parag_range = doc.Paragraphs(i+2).Range
                    parag_range.Font.Name = 'Arial'
                    parag_range.Font.Size = 14
                    parag_range.Font.Bold = constants.wdToggle
                    parag_range.Font.Size = 12
                    doc.TablesOfContents.Add(Range=parag_range,
                                             UseHeadingStyles=True,
                                             UpperHeadingLevel=1,
                                             LowerHeadingLevel=4)
                except Exception as e:
                    print("Ja：", e, "Nein")
                break

    elif toc_count == 1:
        toc = doc.TablesOfContents(1)
        toc.Update()
        print('TOC should have been updated.')
    else:
        print('TOC has not been updated for sure...')    
        
print("Finish AsNiDefFa2.py!")
