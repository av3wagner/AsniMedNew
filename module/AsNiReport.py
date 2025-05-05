#;;*****************************************************************;;;
#;;*****************************************************************;;;
#;;;****************************************************************;;;
#;;;***  FIRMA          : PARADOX                                ***;;;
#;;;***  Autor          : Alexander Wagner                       ***;;;
#;;;***  STUDIEN-NAME   : AsNiFen/Phase II                       ***;;;
#;;;***  STUDIEN-NUMMER :                                        ***;;;
#;;;***  SPONSOR        :                                        ***;;;
#;;;***  ARBEITSBEGIN   : 01.11.2023 / 25.08.2024                ***;;;
#;;;****************************************************************;;;
#;;;*--------------------------------------------------------------*;;;
#;;;*---  PROGRAMM      :AsNiPa203.ipynb                       ---*;;;
#;;;*---  Parent        : AsNiPa202.ipynb, 27.08.2024           ---*;;;
#;;;*---  BESCHREIBUNG  : System                                ---*;;;
#;;;*---                :                                       ---*;;;
#;;;*---                :                                       ---*;;;
#;;;*---  VERSION   VOM : 26.08.2024                            ---*;;;
#;;;*--   KORREKTUR VOM : 04.09.2024                            ---*;;;
#;;;*--                 :                                       ---*;;;
#;;;*---  INPUT         :.INI                                   ---*;;;
#;;;*---  OUTPUT        :                                       ---*;;;
#;;;*--------------------------------------------------------------*;;;
#;;;************************ Änderung ******************************;;;
#;;;****************************************************************;;;
#;;;  Wann              :               Was                        *;;;
#;;;*--------------------------------------------------------------*;;;
#;;;* 04.09.2024        : Komplett Code in einer Cell              *;;;
#;;;* 27.08.2024        : #Funktioniert fehelefrei 10:24 Result:   *;;;
#;;;*                   : #ASNI_ReportResultPa01_20240827OK.docx   *;;;
#;;;****************************************************************;;;

########################## Test0.py #############################
print("Start Step0!") 
#Spire
from spire.doc.common import *
from spire.doc import *

# Используемые стили
HEADER_STYLE = "BoldHeader"
HEADER_LINK_STYLE = "BoldHeaderHyperlink"
CONTENT_STYLE = "Content"
CODE_STYLE = "Code"

document = Document()
section = document.AddSection()
paragraph = section.AddParagraph()
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text =paragraph.AppendText("{{Projekt}}")  
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 26
text.CharacterFormat.Bold = True
text.CharacterFormat.TextColor = Color.get_Blue()

paragraph = section.AddParagraph()
paragraph = section.AddParagraph() 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("Тема исследования: {{Projekt3}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 16
text.CharacterFormat.Bold = True

paragraph = section.AddParagraph()
paragraph = section.AddParagraph() 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("Проект: {{Thema}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 16
text.CharacterFormat.Bold = True

paragraph = section.AddParagraph()
paragraph = section.AddParagraph()  
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("Автор исследования: {{Forscher}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 16
text.CharacterFormat.Bold = True
for num in range(9):
    i=num+1
    paragraph = section.AddParagraph()  
    
paragraph = section.AddParagraph()
paragraph = section.AddParagraph() 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("{{logo}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 16
text.CharacterFormat.Bold = True

for num in range(13):
    i=num+1
    paragraph = section.AddParagraph()  
    
paragraph = section.AddParagraph()  
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("{{Site}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 14
text.CharacterFormat.Bold = True

paragraph = section.AddParagraph() 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("{{Year}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 14
text.CharacterFormat.Bold = True

section = document.AddSection()
paragraph = section.AddParagraph()
paragraph = section.AddParagraph()

paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("Содержание") 
text.CharacterFormat.FontName = "Arial"
text.CharacterFormat.FontSize = 14
text.CharacterFormat.Bold = True
text.CharacterFormat.Italic = True

paragraph = section.AddParagraph()
text = paragraph.AppendText(" ") 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center

outputFile = InEDA_File #"TemplateTOC1.docx"  
document.SaveToFile(outputFile, FileFormat.Docx)
document.Close()

# Используемые стили
HEADER_STYLE = "BoldHeader"
HEADER_LINK_STYLE = "BoldHeaderHyperlink"
CONTENT_STYLE = "Content"
CODE_STYLE = "Code"

document = Document(InEDA_File) 
print("config_json: ", config_json)
def getConfigFile(config):
    with open(config, encoding='utf-8') as json_file:
        return json.load(json_file)
        
config = getConfigFile(config_json)
tasks = config["tasks"]
    
i=-1
k=-1
for task in tasks:
    i=i+1
    header = f"{task['topic']}"
    print("header Heading2: ", i, " ", header)
    Text = f"{task['Text']}"
    Hstyle = f"{task['Header']}"
        
    if header != "Предисловие":
        section = document.AddSection()
        paragraph = section.AddParagraph()
        
    if header != "Содержание":
        paragraph.AppendText(header) 
        if Hstyle == "BuiltinStyle.Heading2":
            paragraph.ApplyStyle(BuiltinStyle.Heading2)
        if Hstyle == "BuiltinStyle.Heading3":
             paragraph.ApplyStyle(BuiltinStyle.Heading3)  
                
    if header == "Предисловие":
        paragraph = section.AddParagraph()
        text = paragraph.AppendText(".") 
        text.CharacterFormat.FontName = "Times New Roman"
        text.CharacterFormat.FontSize = 1
        paragraph.Format.HorizontalAlignment = HorizontalAlignment.Left
        
    if Text != " ": 
            models = task["models"]
            paragraph = section.AddParagraph()
            paragraph.AppendText(Text)
            paragraph.Format.HorizontalAlignment = HorizontalAlignment.Left
 
    if header == "Моделирование": 
        models = task["models"]
        for m in models:
            k=k+1
            Rname = f"Модель: {m}" 
            paragraph = section.AddParagraph()
            paragraph.AppendText(Rname)  
            paragraph.ApplyStyle(BuiltinStyle.Heading4)
                        
            paragraph = section.AddParagraph()
            paragraph.AppendText("ClassBlk" + str(k))  
            paragraph.Format.HorizontalAlignment = HorizontalAlignment.Left

            paragraph = section.AddParagraph()
            paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
            text = paragraph.AppendText("Таблица классификации") 
            text.CharacterFormat.FontName = "Times New Roman"
            text.CharacterFormat.FontSize = 12
            text.CharacterFormat.Bold = True
            paragraph = section.AddParagraph()
            paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
            AddBookmark("t" + str(k), "Table" +  str(k))

            paragraph = section.AddParagraph()     
            paragraph = section.AddParagraph() 
            paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
            text = paragraph.AppendText("Confusion Matrix") 
            text.CharacterFormat.FontName = "Times New Roman"
            text.CharacterFormat.FontSize = 12
            text.CharacterFormat.Bold = True

            paragraph = section.AddParagraph() 
            paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
            paragraph.AppendText("{{Heatmap" + str(k) + "}}")

            paragraph = section.AddParagraph() 
            paragraph = section.AddParagraph() 
            paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center

            text = paragraph.AppendText("ROC Curve") 
            text.CharacterFormat.FontName = "Times New Roman"
            text.CharacterFormat.FontSize = 12
            text.CharacterFormat.Bold = True
            paragraph = section.AddParagraph() 
            paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
            paragraph.AppendText("{{PltR" + str(k) + "}}")

            paragraph = section.AddParagraph() 
            paragraph = section.AddParagraph() 
            paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center

            text = paragraph.AppendText("Score plot") 
            text.CharacterFormat.FontName = "Times New Roman"
            text.CharacterFormat.FontSize = 12
            text.CharacterFormat.Bold = True
            paragraph = section.AddParagraph() 
            paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
            paragraph.AppendText("{{PltL" + str(k) + "}}")

outputFile = "TemplateTOC.docx"        
document.SaveToFile(outputFile, FileFormat.Docx)
document.Close()
print("Programm Test0.py Ende!")

########################## Test1.py #############################
#################################################################
#%matplotlib inline 
print("Start Test1!")
pd.set_option("display.max_rows",None) 

des0=raw_df[raw_df['HeartDisease']==0].describe().T.applymap('{:,.2f}'.format)
des1=raw_df[raw_df['HeartDisease']==1].describe().T.applymap('{:,.2f}'.format)

cat = ['Sex', 'ChestPainType','FastingBS','RestingECG','ExerciseAngina',  'ST_Slope','HeartDisease']
num = ['Age','RestingBP','Cholesterol','MaxHR','Oldpeak']
numerical_columns = []
categorical_columns = []

##########################################################
numerical_columns = list(raw_df.loc[:,['Age', 'RestingBP', 'Cholesterol', 'MaxHR', 'Oldpeak', 'HeartDisease']])
categorical_columns = list(raw_df.loc[:,['Sex', 'ChestPainType', 'FastingBS', 'RestingECG', 'ExerciseAngina', 'ST_Slope']])
numerical=numerical_columns

index = 0
plt.figure(figsize=(20,20))
for feature in numerical:
    if feature != "HeartDisease":
        index += 1
        plt.subplot(2, 3, index)
        sns.boxplot(x='HeartDisease', y=feature, data=raw_df)
        
plt.savefig(pathIm + '/EDA1.png')  

print('numerical_columns before clear:', numerical_columns)
numerical_columns.clear()
print('numerical_columns after clear:', numerical_columns)
del numerical_columns[:]
del categorical_columns[:]
numerical_columns = list(raw_df.loc[:,['Age', 'RestingBP', 'Cholesterol', 'MaxHR', 'Oldpeak', 'HeartDisease']])
categorical_columns = list(raw_df.loc[:,['Sex', 'ChestPainType', 'FastingBS', 'RestingECG', 'ExerciseAngina', 'ST_Slope']])

# checking boxplots

boxplots_custom(dataset=raw_df, columns_list=numerical_columns, rows=2, cols=3, suptitle='Boxplots for each variable')
plt.tight_layout()
plt.savefig(pathIm + '/EDA2.png')

fig, axes = plt.subplots(nrows=3, ncols=2,figsize=(11,17))
fig.suptitle('Features vs Class\n', size = 18)

sns.boxplot(ax=axes[0, 0], data=raw_df, x='Sex', y='Age', palette='Spectral')
axes[0,0].set_title("Age distribution");


sns.boxplot(ax=axes[0,1], data=raw_df, x='Sex', y='RestingBP', palette='Spectral')
axes[0,1].set_title("RestingBP distribution");


sns.boxplot(ax=axes[1, 0], data=raw_df, x='Sex', y='Cholesterol', palette='Spectral')
axes[1,0].set_title("Cholesterol distribution");

sns.boxplot(ax=axes[1, 1], data=raw_df, x='Sex', y='MaxHR', palette='Spectral')
axes[1,1].set_title("MaxHR distribution");

sns.boxplot(ax=axes[2, 0], data=raw_df, x='Sex', y='Oldpeak', palette='Spectral')
axes[2,0].set_title("Oldpeak distribution");

sns.boxplot(ax=axes[2, 1], data=raw_df, x='Sex', y='HeartDisease', palette='Spectral')
axes[2,1].set_title("HeartDisease distribution");

plt.tight_layout()
plt.savefig(pathIm + '/EDA3.png')

######################## numeric_columns ################################
numeric_columns = list(raw_df.loc[:,['Age', 'RestingBP', 'Cholesterol', 'MaxHR', 'Oldpeak', 'HeartDisease']])

fig = plt.figure(figsize=(15, 10))
plt.title('Kdeplot для цифровых переменных, категория: HeartDisease')
rows, cols = 2, 3
for idx, num in enumerate(numeric_columns[:30]):
    ax = fig.add_subplot(rows, cols, idx+1)
    ax.grid(alpha = 0.7, axis ="both")
    sns.kdeplot(x = num, fill = True,color ="#3386FF",linewidth=0.6, data = raw_df[raw_df['HeartDisease']==0], label = "Healthy")
    sns.kdeplot(x = num, fill = True,color ="#EFB000",linewidth=0.6, data = raw_df[raw_df['HeartDisease']==1] , label = "Heart Disease")
    ax.set_xlabel(num)
    ax.legend()
    
fig.tight_layout()
plt.savefig(pathIm + '/EDA4.png')

fig = plt.figure(figsize=(15, 10))
plt.title('Kdeplot для цифровых переменных, категория: Sex')
rows, cols = 2, 3
for idx, num in enumerate(numeric_columns[:30]):
    ax = fig.add_subplot(rows, cols, idx+1)
    ax.grid(alpha = 0.7, axis ="both")
    sns.kdeplot(x = num, fill = True,color ="#3386FF",linewidth=0.6, data = raw_df[raw_df['Sex']=="M"], label = "M")
    sns.kdeplot(x = num, fill = True,color ="#EFB000",linewidth=0.6, data = raw_df[raw_df['Sex']=="F"], label = "F")
    ax.set_xlabel(num)
    ax.legend()
    
fig.tight_layout()
plt.savefig(pathIm + '/EDA5.png')

raw_df = pd.read_csv(pfad+'/heart.csv')
fig = plt.figure(figsize=(25, 10))
fig = px.scatter_3d(raw_df, 
                    x='RestingBP',
                    y='Age',
                    z='Sex',
                    color='HeartDisease')

fig.write_html(pathIm + '/Buble3D.html')

with open(pathIm + "/EDA6.png", 'wb') as f:
    f.write(pplt.io.to_image(fig, width=1200, height=800, format='png'))   

df = pd.read_csv(pfad+'/heart.csv')
replace_zero_cholesterol(df)

fig = px.scatter(df, y = 'Age',x='Cholesterol', color='Cholesterol' )
fig.update_layout(title=f'Buble Chart Cholesterol')
pio.write_image(fig, pathIm + '/EDA7.png', width=1200, height=800, format='png', scale=6)

cf.go_offline()
cf.set_config_file(offline=True, world_readable=True)
warnings.filterwarnings('ignore')
warnings.warn("this will not show")
plt.rcParams["figure.figsize"] = (10,6)
pd.set_option('max_colwidth',200)
pd.set_option('display.max_columns', 200)
pd.set_option('display.float_format', lambda x: '%.3f' % x)

cat = ['Sex', 'ChestPainType','FastingBS','RestingECG',
                          'ExerciseAngina',  'ST_Slope','HeartDisease']
num = ['Age','RestingBP','Cholesterol','MaxHR','Oldpeak']

df=raw_df
fig = px.scatter(df, 
                 x=df.Age, 
                 y=df.Cholesterol, 
                 color=df.HeartDisease, 
                 facet_col=df.FastingBS,
                 facet_row=df.Sex,
                 color_discrete_map={1: "#FF5722",0: "#7CB342"},
                 width=950, 
                 height=800,
                 title="HeartDisease Data")

fig.update_layout(
                    plot_bgcolor= "#dcedc1",
                    paper_bgcolor="#FFFDE7",
                 )

fig.write_image(pathIm + '/EDA11.png',scale=4)

colors = px.colors.cyclical.Twilight

HDValues={
    0:'Healthy',
    1:'Heart Disease'
    }

df = raw_df
sns.set_theme(rc = {'figure.dpi': 250, 'axes.labelsize': 7, 
                    'axes.facecolor': '#FFFDE7', 'grid.color': '#fffdfa', 
                    'figure.facecolor': '#FFFDE7'}, font_scale = 0.55)
fig, ax = plt.subplots(4, 2, figsize = (6.5, 7.5))
for indx, (column, axes) in list(enumerate(list(zip(cat, ax.flatten())))):
    
    sns.violinplot(ax = axes, x = df[column], 
                   y = df['Age'],
                   scale = 'width', linewidth = 0.5, 
                   palette = colors, inner = None)
    
    plt.setp(axes.collections, alpha = 0.3)
    
    sns.stripplot(ax = axes, x = df[column], 
                  y = df['Age'],
                  palette = colors, alpha = 0.9, 
                  s = 1.5, jitter = 0.07)
    sns.pointplot(ax = axes, x = df[column],
                  y = df['Age'],
                  color = '#ff5736', scale = 0.25,
                  estimator = np.mean, ci = 'sd',
                  errwidth = 0.5, capsize = 0.15, join = True)
    
    plt.setp(axes.lines, zorder = 100)
    plt.setp(axes.collections, zorder = 100)
    
else:
    [axes.set_visible(False) for axes in ax.flatten()[indx + 1:]]
    
plt.tight_layout()
fig.savefig(pathIm + '/EDA12.png')


sns.set_theme(rc = {'figure.dpi': 120, 'axes.labelsize': 8, 
                    'axes.facecolor': '#FFFDE7', 'grid.color': '#fffdfa', 
                    'figure.facecolor': '#FFFDE7'}, font_scale = 0.65)

fig, ax = plt.subplots(5, 1, figsize = (10, 10))

for indx, (column, axes) in list(enumerate(list(zip(num, ax.flatten())))):
    
    sns.scatterplot(ax = axes, y = df[column].index, x = df[column], 
                    hue = df['HeartDisease'], palette = 'magma', alpha = 0.8)
    
else:
    [axes.set_visible(False) for axes in ax.flatten()[indx + 1:]]
    
plt.tight_layout()
fig.savefig(pathIm + '/EDA13.png')

sns.set_theme(rc = {'figure.dpi': 120, 'axes.labelsize': 8, 
                    'axes.facecolor': '#FFFDE7', 'grid.color': '#fffdfa', 
                    'figure.facecolor': '#FFFDE7'}, font_scale = 0.65)

fig, ax = plt.subplots(5, 1, figsize = (10, 14))

for indx, (column, axes) in list(enumerate(list(zip(num, ax.flatten())))):
    
    sns.histplot(ax = axes, x = df[column], hue = df['HeartDisease'], 
                 palette = 'magma', alpha = 0.8, multiple = 'stack')
    
    legend = axes.get_legend() # sns.hisplot has some issues with legend
    handles = legend.legendHandles
    legend.remove()
    axes.legend(handles, ['0', '1'], title = 'HeartDisease', loc = 'upper right')
    Quantiles = np.quantile(df[column], [0, 0.25, 0.50, 0.75, 1])
    
    for q in Quantiles: axes.axvline(x = q, linewidth = 0.5, color = 'r')
        
plt.tight_layout()
fig.savefig(pathIm + '/EDA14.png')

df2 = df.groupby('Sex').agg({'Age' : 'mean', "ChestPainType":'count','RestingBP':'mean','Cholesterol':'mean',
                            'FastingBS':'sum','RestingECG':'count','MaxHR':'mean','ExerciseAngina':'count','Oldpeak':'mean',
                            'ST_Slope':'count','HeartDisease':'sum'})
df2
fig=px.bar(data_frame=df2, barmode='group',
       title = "<b>Gender wise Analyzing</b>",template="plotly_dark")
fig.write_image(pathIm + '/EDA15.png',scale=4)

try:
    heart_dft = pd.read_csv(pfad+'/heart.csv')
except:
    heart_dft=pd.read_csv(pfad+'/heart.csv')

sex_color = dict({"Male": "#2986cc", "Female": "#c90076"})
plt.style.use("fivethirtyeight")
heart_dft["Sex"] = heart_dft["Sex"].map({"M": "Male", "F": "Female"})
heart_dft["Sex"]
heart_dft["HeartDisease"] = heart_dft["HeartDisease"].map({0: "No", 1: "Yes"})

filtheart_dft = heart_dft["Cholesterol"] > 0
heart_dft_chol_n0 = heart_dft[filtheart_dft]

#print(heart_dft.head())

sex_color = dict({"Male": "#2986cc", "Female": "#c90076"})
plt.style.use("fivethirtyeight")
heart_dft["Sex"] = heart_dft["Sex"].map({"M": "Male", "F": "Female"})
heart_dft["Sex"]
heart_dft["HeartDisease"] = heart_dft["HeartDisease"].map({0: "No", 1: "Yes"})
filtheart_dft = heart_dft["Cholesterol"] > 0
heart_dft_chol_n0 = heart_dft[filtheart_dft]

g=sns.JointGrid(
    data=heart_dft, x="Age", y="Cholesterol", hue="Sex", palette=sex_color
).plot(sns.scatterplot, sns.histplot)

plt.legend(title='Company', fontsize=20)
plt.xlabel('Agex', fontsize=10);
plt.ylabel('Cholesterolx', fontsize=10);
plt.title('Sales Data', fontsize=12)
plt.tick_params(axis='both', which='major', labelsize=10)
g.savefig(pathIm + '/EDA16.png')
sns.set_theme()
pw.overwrite_axisgrid() 
iris = sns.load_dataset("iris")
tips = sns.load_dataset("tips")

# An lmplot
g0 = sns.lmplot(x="total_bill", y="tip", hue="smoker", data=tips, 
                palette=dict(Yes="g", No="m"))
g0 = pw.load_seaborngrid(g0, label="g0")

# A Pairplot
g1 = sns.pairplot(iris, hue="species")
g1 = pw.load_seaborngrid(g1, label="g1")

# A relplot
g2 = sns.relplot(data=tips, x="total_bill", y="tip", col="time", hue="time", 
                 size="size", style="sex", palette=["b", "r"], sizes=(10, 100))
g2 = pw.load_seaborngrid(g2, label="g2")

g3 = sns.jointplot(x='Cholesterol',y='Age',data=raw_df, hue="Sex") 
g3 = pw.load_seaborngrid(g3, label="g3")
(((g0|g1)["g0"]/g3)["g3"]|g2).savefig(pathIm + '/EDA17.png')

try:
    heart_dft = pd.read_csv(pfad+'/heart.csv')
except:
    heart_dft=pd.read_csv(pfad+'/heart.csv')


sex_color = dict({"Male": "#2986cc", "Female": "#c90076"})
plt.style.use("fivethirtyeight")
heart_dft["Sex"] = heart_dft["Sex"].map({"M": "Male", "F": "Female"})
heart_dft["Sex"]
heart_dft["HeartDisease"] = heart_dft["HeartDisease"].map({0: "No", 1: "Yes"})

filtheart_dft = heart_dft["Cholesterol"] > 0
heart_dft_chol_n0 = heart_dft[filtheart_dft]

Chol_mean_f = (
    heart_dft_chol_n0[["Sex", "Cholesterol"]]
    .groupby(["Sex"])
    .mean("Cholesterol")
    .loc["Female", "Cholesterol"]
).round()

Chol_mean_m = (
    heart_dft_chol_n0[["Sex", "Cholesterol"]]
    .groupby(["Sex"])
    .mean("Cholesterol")
    .loc["Male", "Cholesterol"]
).round()

plt.figure(figsize=(10, 5))
sns.set_context("paper")

kdeplt = sns.kdeplot(
    data=heart_dft_chol_n0,
    x="Cholesterol",
    hue="Sex",
    palette=sex_color,
    alpha=0.7,
    lw=2,
)

kdeplt.set_title("Cholesterol values distribution\n Male VS Female", fontsize=12)
kdeplt.set_xlabel("Cholesterol", fontsize=12)
plt.axvline(x=Chol_mean_f, color="#c90076", ls="--", lw=1.3)
plt.axvline(x=Chol_mean_m, color="#2986cc", ls="--", lw=1.3)
plt.text(108, 0.00612, "Mean Cholesterol / Male", fontsize=10, color="#2986cc")
plt.text(260, 0.006, "Mean Cholesterol / Female", fontsize=10, color="#c90076")
kdeplt.figure.savefig(pathIm + '/EDA18.png')

#######################################################
################### SAS GRAPH 19-23 ###################
#######################################################

#Kategorial
pio.renderers
def auto_fmt (pct_value):
    return '{:.0f}\n({:.1f}%)'.format(raw_df['HeartDisease'].value_counts().sum()*pct_value/100,pct_value) 

try:
    raw_df = pd.read_csv(pfad+'/heart.csv')
except:
    raw_df = pd.read_csv(pfad+'/heart.csv')
HDValues={
    0:'Healthy',
    1:'Heart Disease'
    }

df = raw_df.HeartDisease.replace(HDValues)

fig=plt.figure(figsize=(6, 6))
matplotlib.rcParams.update({'font.size': 15})

df.value_counts().plot.pie(explode=[0.1, 0.1],                               
                                       autopct=auto_fmt,
                                       textprops={'fontsize': 16},
                                       shadow=True)

plt.title('Healthy vs Heart Disease', color='Red',pad=15, fontsize=20);
plt.axis('off');
plt.savefig(pathIm + '/EDA31.png')

plt.figure(figsize=(6, 6))
matplotlib.rcParams.update({'font.size': 15})

raw_df.Sex.value_counts().plot.pie(explode=[0.1, 0.1],
                                       #autopct='%1.2f%%',
                                       autopct=auto_fmt,
                                       textprops={'fontsize': 16},
                                       shadow=True)
plt.title('Sex', color='Red',pad=10, fontsize=20);
plt.axis('off');
plt.savefig(pathIm + '/EDA32.png')

"""
categorical = list(raw_df.loc[:,['Sex', 'ChestPainType', 'FastingBS', 'RestingECG', 'ExerciseAngina', 'ST_Slope']])
raw_df[categorical].iplot(kind='box', subplots=True,bins=50, theme='white', asImage=True, dimensions=(800,500))    
fig=raw_df[categorical].iplot(kind='box', subplots=True, bins=50, theme='white', asFigure=True, dimensions=(800,500))    
fig.write_image(pathIm + "/EDA34.png")

df1=raw_df[categorical]
df1.iplot(kind='hist', theme='white',asImage=True,dimensions=(1200,800))   
fig=df1.iplot(kind='hist', theme='white',asFigure=True, dimensions=(1200,800))   
fig.write_image(pathIm + "/EDA35.png")
"""

fig, ax = plt.subplots (3, 2, figsize=(16, 16))
ax_rst = []
for i in range(len(categorical_columns)):
    axs = sns.countplot(data=raw_df, x =raw_df[categorical_columns[i]], ax=ax[int(i/2),i % 2])
    ax_rst.append(axs)
    total = raw_df[categorical_columns[i]].value_counts().sum()
    for p in axs.patches:
        value_pct = '{:.0f} ({:.1f}%)'.format(p.get_height(), 100 * p.get_height()/total)
        x = p.get_x() + p.get_width()/2
        y = p.get_height()
        axs.annotate(value_pct, (x, y),ha='center')   
plt.savefig(pathIm + '/EDA33.png')

fig=px.pie(raw_df,values='HeartDisease',names='ChestPainType', 
           template='plotly_dark',color_discrete_sequence=px.colors.sequential.RdBu,
           title='The effect of the type of chest pain on the disease')
fig.update_traces(textposition='inside',textinfo='percent+label')
fig.update_layout(width=1000, height=800)
fig.write_image(pathIm + '/EDA36.png', scale=4)

fig=px.pie(raw_df,values='HeartDisease',names='ST_Slope',hole=.4,template='plotly_dark',title='The effect of the the slope of the peak exercise on the disease',)
fig.update_traces(textposition='inside',textinfo='percent+label')
fig.update_layout(annotations=[dict(text='ST slope', x=0.5, y=0.5, font_size=20, showarrow=False)])
fig.update_layout(width=1000, height=1000)
fig.write_image(pathIm + '/EDA37.png',scale=4)

df=raw_df
colors = px.colors.cyclical.Twilight
fig = make_subplots(rows=1,cols=2,
                    subplot_titles=('Countplot',
                                    'Percentages'),
                    specs=[[{"type": "xy"},
                            {'type':'domain'}]])

fig.add_trace(go.Bar(y = df['Sex'].value_counts().values.tolist(), 
                      x = df['Sex'].value_counts().index, 
                      text=df['Sex'].value_counts().values.tolist(),
              textfont=dict(size=15),
                      textposition = 'outside',
                      showlegend=False,
              marker = dict(color = colors,
                            line_color = 'black',
                            line_width=3)),row = 1,col = 1)
fig.add_trace((go.Pie(labels=df['Sex'].value_counts().keys(),
                             values=df['Sex'].value_counts().values,textfont = dict(size = 16),
                     hole = .4,
                     marker=dict(colors=colors),
                     textinfo='label+percent',
                     hoverinfo='label')), row = 1, col = 2)
fig.update_yaxes(range=[0,800])
fig.update_layout(
                    paper_bgcolor= '#FFFDE7',
                    plot_bgcolor= '#FFFDE7',
                    title=dict(text = "Gender Distribution",x=0.5,y=0.95),
                    title_font_size=30
                  )
iplot(fig)
fig.write_image(pathIm + '/EDA38.png',scale=4)

colors = px.colors.cyclical.Twilight
fig = make_subplots(rows=1,cols=2,
                    subplot_titles=('Countplot',
                                    'Percentages'),
                    specs=[[{"type": "xy"},
                            {'type':'domain'}]])
fig.add_trace(go.Bar(y = df['HeartDisease'].value_counts().values.tolist(), 
                      x = df['HeartDisease'].value_counts().index, 
                      text=df['HeartDisease'].value_counts().values.tolist(),
              textfont=dict(size=15),
                      textposition = 'outside',
                      showlegend=False,
              marker = dict(color = colors,
                            line_color = 'black',
                            line_width=3)),row = 1,col = 1)
fig.add_trace((go.Pie(labels=df['HeartDisease'].value_counts().keys(),
                             values=df['HeartDisease'].value_counts().values,textfont = dict(size = 16),
                     hole = .4,
                     marker=dict(colors=colors),
                     textinfo='label+percent',
                     hoverinfo='label')), row = 1, col = 2)
fig.update_yaxes(range=[0,550])
fig.update_layout(
                    paper_bgcolor= '#FFFDE7',
                    plot_bgcolor= '#FFFDE7',
                    title=dict(text = "HeartDisease Distribution",x=0.5,y=0.95),
                    title_font_size=30
                  )
iplot(fig)
fig.write_image(pathIm + '/EDA39.png',scale=4)  

cat = ['Sex', 'ChestPainType','FastingBS','RestingECG',
                          'ExerciseAngina',  'ST_Slope','HeartDisease']
num = ['Age','RestingBP','Cholesterol','MaxHR','Oldpeak']

import seaborn as sns
sns.set_theme(rc = {'figure.dpi': 250, 'axes.labelsize': 7, 
                    'axes.facecolor': '#FFFDE7', 'grid.color': '#fffdfa', 
                    'figure.facecolor': '#FFFDE7'}, font_scale = 0.55)
fig, ax = plt.subplots(3, 2, figsize = (6.5, 9))
for indx, (column, axes) in list(enumerate(list(zip(cat, ax.flatten())))):
    if column not in 'HearDisease':
        sns.countplot(ax = axes, x = df[column], hue = df['HeartDisease'], palette = colors, alpha = 1)  
else:
    [axes.set_visible(False) for axes in ax.flatten()[indx + 1:]]   
    
axes_legend = ax.flatten()
axes_legend[1].legend(title = 'HeartDisease', loc = 'upper right')
axes_legend[2].legend(title = 'HeartDisease', loc = 'upper right')
fig.savefig(pathIm + '/EDA40.png')

#import seaborn as sns
sns.set_theme(rc = {'figure.dpi': 250, 'axes.labelsize': 7, 
                    'axes.facecolor': '#FFFDE7', 'grid.color': '#fffdfa', 
                    'figure.facecolor': '#FFFDE7'}, font_scale = 0.55)
fig, ax = plt.subplots(3, 2, figsize = (6.5, 9))
for indx, (column, axes) in list(enumerate(list(zip(cat[1:], ax.flatten())))):
    sns.countplot(ax = axes, x = df[column], hue = df['Sex'], palette = colors, alpha = 1)  
else:
    [axes.set_visible(False) for axes in ax.flatten()[indx + 1:]]   
axes_legend = ax.flatten()
axes_legend[1].legend(title = 'Sex', loc = 'upper right')
axes_legend[2].legend(title = 'Sex', loc = 'upper right')
fig.savefig(pathIm + '/EDA41.png')

sns.set_theme(rc = {'figure.dpi': 250, 'axes.labelsize': 7, 
                    'axes.facecolor': '#FFFDE7', 'grid.color': '#fffdfa', 
                    'figure.facecolor': '#FFFDE7'}, font_scale = 0.55)
fig, ax = plt.subplots(3, 2, figsize = (6.5, 9))
cat2 = []
for i in cat:
    if i not in 'ChestPainType':
        cat2.append(i)
for indx, (column, axes) in list(enumerate(list(zip(cat2, ax.flatten())))):
    sns.countplot(ax = axes, x = df[column], hue = df['ChestPainType'], palette = colors, alpha = 1)  
else:
    [axes.set_visible(False) for axes in ax.flatten()[indx + 1:]]   
axes_legend = ax.flatten()
axes_legend[1].legend(title = 'ChestPainType', loc = 'upper right')
axes_legend[2].legend(title = 'ChestPainType', loc = 'upper right')
fig.savefig(pathIm + '/EDA42.png')

fig, ax = plt.subplots() 
ax = plt.subplot(1,2,1)
ax = sns.countplot(x='Sex', data=raw_df)
ax.bar_label(ax.containers[0])
ax =plt.subplot(1,2,2)
ax=raw_df['Sex'].value_counts().plot.pie(explode=[0.1, 0.1],autopct='%1.2f%%',shadow=True);
ax.set_title(label = "Sex", fontsize = 16) #,color='Red',font='Lucida Calligraphy')
plt.savefig(pathIm + '/EDA43.png')
fig.clear(True)

fig, ax1 = plt.subplots()
heart=raw_df
ax1 = plt.subplot(1,2,1)
ax1 = sns.countplot(x='ChestPainType', data=heart)
ax1.bar_label(ax1.containers[0])
plt.title("ChestPainType", fontsize=14)
ax1 =plt.subplot(1,2,2)
ax1=heart['ChestPainType'].value_counts().plot.pie(explode=[0.1, 0.1,0.1,0.1],autopct='%1.2f%%',shadow=True);
ax1.set_title(label = "ChestPainType", fontsize = 20,color='Red',font='Lucida Calligraphy');
plt.savefig(pathIm + '/EDA44.png')
fig.clear(True)

fig, ax2 = plt.subplots()
ax2 = plt.subplot(1,2,1)
ax2 = sns.countplot(x='RestingECG', data=heart)
ax2.bar_label(ax2.containers[0])
plt.title("RestingECG", fontsize=14)

ax2 =plt.subplot(1,2,2)
ax2=heart['RestingECG'].value_counts().plot.pie(explode=[0.1, 0.1,0.1],autopct='%1.2f%%',shadow=True);
ax2.set_title(label = "RestingECG", fontsize = 20,color='Red',font='Lucida Calligraphy');
plt.savefig(pathIm + '/EDA45.png')
time.sleep(1)
fig.clear(True)

fig, ax3 = plt.subplots()
ax3 = plt.subplot(1,2,1)
ax3 = sns.countplot(x='ExerciseAngina', data=heart)
ax3.bar_label(ax3.containers[0])
plt.title("ExerciseAngina", fontsize=14)

ax3 =plt.subplot(1,2,2)
ax3=heart['ExerciseAngina'].value_counts().plot.pie(explode=[0.1, 0.1],autopct='%1.2f%%',shadow=True);
ax3.set_title(label = "ExerciseAngina", fontsize = 20,color='Red',font='Lucida Calligraphy');
plt.savefig(pathIm + '/EDA46.png')
fig.clear(True)

fig, ax = plt.subplots()
ax = plt.subplot(1,2,1)
ax = sns.countplot(x='ST_Slope', data=heart)
ax.bar_label(ax.containers[0])
plt.title("ST_Slope", fontsize=14)

ax =plt.subplot(1,2,2)
ax=heart['ST_Slope'].value_counts().plot.pie(explode=[0.1, 0.1,0.1],autopct='%1.2f%%',shadow=True);
ax.set_title(label = "ST_Slope", fontsize = 20,color='Red',font='Lucida Calligraphy');
plt.savefig(pathIm + '/EDA47.png')
fig.clear(True)

#heart=pd.read_csv('data/heart.csv')
sns.set(font_scale=1.1)
heart["Cholesterol_Category"]= pd.cut(heart["Cholesterol"] ,bins=[0, 200, 230 , 500] ,labels=["Normal","Borderline","High" ] )
print("Value Counts :\n\n",heart['Cholesterol_Category'].value_counts())

heart.head()
fig, ax = plt.subplots()
ax = plt.subplot(1,2,1)
ax = sns.countplot(x='Cholesterol_Category', data=heart)
ax.bar_label(ax.containers[0])
plt.title("Cholesterol_Categoryy", fontsize=14)

ax =plt.subplot(1,2,2)
ax=heart['Cholesterol_Category'].value_counts().plot.pie(explode=[0.1, 0.1,0.1],autopct='%1.2f%%',shadow=True);
ax.set_title(label = "Cholesterol_Category", fontsize = 20,color='Red',font='Lucida Calligraphy');
plt.savefig(pathIm + '/EDA48.png')
fig.clear(True)

heart["RestingBP_Category"]= pd.cut(heart["RestingBP"] ,bins=[0,120, 129 , 139,200] ,labels=["Normal_BP","Elevated_BP","Hypertension_Stage_1", "Hypertension_Stage_2"] )
print("Value Counts :\n\n",heart['RestingBP_Category'].value_counts())
heart.sample(5)
heart['RestingBP_Category'] = heart['RestingBP_Category'].astype(object)

plt.rcParams['legend.fontsize'] = 7
sns.set(font_scale=1.0)
fig, ax = plt.subplots()
ax = plt.subplot(1,2,1)
ax = sns.countplot(x='RestingBP_Category', data=heart)
ax.bar_label(ax.containers[0])
plt.axis('off');

ax =plt.subplot(1,2,2)
ax=heart['RestingBP_Category'].value_counts().plot.pie(explode=[0.1, 0.1,0.1,0.1],autopct='%1.2f%%',shadow=True);
plt.axis('off');
plt.savefig(pathIm + '/EDA49.png')
fig.clear(True)

df = heart  
male_df = df[df['Sex'] == 'M']
female_df = df[df['Sex'] == 'F']

## Grouping Datasets
male_cp_fbs = male_df.groupby(['ChestPainType', 'FastingBS']).size().reset_index().rename(columns={0: 'count'})
female_cp_fbs = female_df.groupby(['ChestPainType', 'FastingBS']).size().reset_index().rename(columns={0: 'count'})

male_st_ecg = male_df.groupby(['ST_Slope', 'RestingECG']).size().reset_index().rename(columns={0: 'count'})
female_st_ecg = female_df.groupby(['ST_Slope', 'RestingECG']).size().reset_index().rename(columns={0: 'count'})

male_ea_cp = male_df.groupby(['ExerciseAngina', 'ChestPainType']).size().reset_index().rename(columns={0: 'count'})
female_ea_cp = female_df.groupby(['ExerciseAngina', 'ChestPainType']).size().reset_index().rename(columns={0: 'count'})

## Creating Sunburst Figures
sb1 = px.sunburst(male_cp_fbs, values='count', path=['ChestPainType', 'FastingBS'])
sb2 = px.sunburst(female_cp_fbs, values='count', path=['ChestPainType', 'FastingBS'])

sb3 = px.sunburst(male_st_ecg, values='count', path=['ST_Slope', 'RestingECG'])
sb4 = px.sunburst(female_st_ecg, values='count', path=['ST_Slope', 'RestingECG'])

sb5 = px.sunburst(male_ea_cp, values='count', path=['ExerciseAngina', 'ChestPainType'])
sb6 = px.sunburst(female_ea_cp, values='count', path=['ExerciseAngina', 'ChestPainType'])

## Subplots
fig = make_subplots(rows=3, cols=2, specs=[
    [{"type": "sunburst"}, {"type": "sunburst"}],
    [{"type": "sunburst"}, {"type": "sunburst"}],
    [{"type": "sunburst"}, {"type": "sunburst"}]
], subplot_titles=("Male Chest Pain with Fasting Blood Sugar", "Female Chest Pain with Fasting Blood Sugar",
                   "Male ST Slope with Resting ECG", "Female ST Slope with Resting ECG",
                   "Male Exercise Angina with Chest Pain Type", "Female Exercise Angina with Chest Pain Type"))

## Plotting Figures
fig.add_trace(sb1.data[0], row=1, col=1)
fig.add_trace(sb2.data[0], row=1, col=2)
fig.add_trace(sb3.data[0], row=2, col=1)
fig.add_trace(sb4.data[0], row=2, col=2)
fig.add_trace(sb5.data[0], row=3, col=1)
fig.add_trace(sb6.data[0], row=3, col=2)

fig.update_traces(textinfo="label+percent parent")
fig.update_layout(title_text="Male vs Female Sunburst", title_x=0.5, 
                  height=1200, width=1200, template='plotly_dark', showlegend=False,
        font=dict(
            family="Rubik",
            size=14)
)

fig.write_image(pathIm + '/EDA50.png',scale=6)

################
heart_dft= heart 
RestingECG_vs_Sex = (
    heart_dft[["RestingECG", "Sex"]]
    .value_counts(normalize=True)
    .reset_index(name="Pct")
    .sort_values(by="RestingECG")
)
RestingECG_vs_Sex["Pct"] = RestingECG_vs_Sex["Pct"].round(2) * 100
RestingECG_vs_Sex.sort_values(by="Pct", ascending=False)

ChestPainType_vs_Sex = (
    heart_dft[["ChestPainType", "Sex"]]
    .value_counts(normalize=True)
    .reset_index(name="Pct")
    .sort_values(by="ChestPainType")
)
ChestPainType_vs_Sex["Pct"] = ChestPainType_vs_Sex["Pct"].round(2) * 100
plt.style.use("fivethirtyeight")
fig, ax = plt.subplots(1, 2, figsize=(12, 5), sharey=True)
palette4 = {"ASY": "#1b85b8", "ATA": "#5a5255", "NAP": "#559e83", "TA": "#ae5a41"}
palette5 = {"LVH": "#2dc937", "Normal": "#e7b416", "ST": "#cc3232"}

sns.barplot(
    data=ChestPainType_vs_Sex,
    x="Sex",
    hue="ChestPainType",
    #errorbar=None,
    y="Pct",
    palette=palette4,
    linewidth=0.5,
    edgecolor="black",
    alpha=0.8,
    ax=ax[0],
)

for ax1 in [ax[0]]:
    for container in ax1.containers:
        values2 = container.datavalues
        labels = ["{:g}%".format(val) for val in values2]
        ax1.bar_label(container, labels=labels)

ax[0].set_ylabel("Percent")
ax[0].set_xlabel("")
ax[0].set_title(
    "Regardless of the proportion of Males and Females,\n Men have high ASY compared with Women, and the pattern is different.",
    fontsize=10,
)

sns.barplot(
    data=RestingECG_vs_Sex,
    x="Sex",
    hue="RestingECG",
    #errorbar=None,
    y="Pct",
    palette=palette5,
    linewidth=0.5,
    edgecolor="black",
    alpha=0.8,
    ax=ax[1],
)

for ax2 in [ax[1]]:
    for container in ax2.containers:
        values3 = container.datavalues
        labels = ["{:g}%".format(val) for val in values3]
        ax2.bar_label(container, labels=labels)

ax[1].set_ylabel("")
ax[1].set_xlabel("")
ax[1].set_title("Men and Women have somehow same pattern of RestingECG", fontsize=10)
plt.tight_layout()
fig.savefig(pathIm + '/EDA51.png')
fig.clear(True)

ExerciseAngina_vs_Sex = (
    heart_dft[["ExerciseAngina", "Sex"]]
    .value_counts(normalize=True)
    .reset_index(name="Pct")
    .sort_values(by="ExerciseAngina")
)
ExerciseAngina_vs_Sex["Pct"] = ExerciseAngina_vs_Sex["Pct"].round(2) * 100


ST_Slope_vs_Sex = (
    heart_dft[["ST_Slope", "Sex"]]
    .value_counts(normalize=True)
    .reset_index(name="Pct")
    .sort_values(by="ST_Slope")
)
ST_Slope_vs_Sex["Pct"] = ST_Slope_vs_Sex["Pct"].round(2) * 100

plt.style.use("fivethirtyeight")
fig, ax = plt.subplots(1, 2, figsize=(12, 5), sharey=True)

palette6 = {
    "Y": "#000000",
    "N": "#009900",
}

palette7 = {"Down": "#b2d8d8", "Flat": "#66b2b2", "Up": "#004c4c"}

sns.barplot(
    data=ExerciseAngina_vs_Sex,
    x="Sex",
    hue="ExerciseAngina",
    #errorbar=None,
    y="Pct",
    palette=palette6,
    linewidth=0.5,
    edgecolor="black",
    alpha=0.8,
    ax=ax[0],
)

for ax3 in [ax[0]]:
    for container in ax3.containers:
        values2 = container.datavalues
        labels = ["{:g}%".format(val) for val in values2]
        ax3.bar_label(container, labels=labels)

ax[0].set_ylabel("Percent")
ax[0].set_xlabel("")
ax[0].set_title(
    "Almost a similar pattern between Men and Women. (ExerciseAngina)", fontsize=10
)

sns.barplot(
    data=ST_Slope_vs_Sex,
    x="Sex",
    hue="ST_Slope",
    #errorbar=None,
    y="Pct",
    palette=palette7,
    linewidth=0.5,
    edgecolor="black",
    alpha=0.8,
    ax=ax[1],
)

for ax4 in [ax[1]]:
    for container in ax4.containers:
        values3 = container.datavalues
        labels = ["{:g}%".format(val) for val in values3]
        ax4.bar_label(container, labels=labels)

ax[1].set_ylabel("")
ax[1].set_xlabel("")
ax[1].set_title(
    "A different pattern between Men and Women (ExerciseAngina)", fontsize=10
)
plt.tight_layout()
fig.savefig(pathIm + '/EDA52.png')
print("Programm Test1.py Ende!")
####################### Test2.py ################################
#################################################################

print("Start Step2!")
today = datetime.date.today()
year = today.year
start_time = time.time()

word_app = win32com.client.Dispatch("Word.Application")
word_app.Visible = True
doc = word_app.Documents.Open(InEDA)
print(len(doc.Paragraphs))
start = []
stop = []
to_start = "xStart"
to_stop = "xStop"

for i, p in enumerate(doc.Paragraphs):
    if to_start in p.Range.Text:
        start.extend([i+1])
    if to_stop in p.Range.Text:
        stop.extend([i+1])    
print(len(start))

# Используемые стили
HEADER_STYLE = "BoldHeader"
HEADER_LINK_STYLE = "BoldHeaderHyperlink"
CONTENT_STYLE = "Content"
CODE_STYLE = "Code"

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

doc = Document() 
InDoc = Document(InEDA) 
config = getConfigFile(eda_json) 

# Deleting heading latent styles seems to do nothing:
latent_styles = doc.styles.latent_styles
latent_styles['Heading 1'].delete()
latent_styles['Heading 2'].delete()
latent_styles['Heading 3'].delete()
styles =  doc.styles

new_heading_style = styles.add_style('New Heading1', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Heading 1']
font = new_heading_style.font
font.name = "Times New Roman"
font.size = Pt(16)
font.bold = True
font.color.rgb = RGBColor(0x00, 0x00, 0x00)

new_heading_style = styles.add_style('New Heading2', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Heading 2']
font = new_heading_style.font
font.name = "Times New Roman"
font.size = Pt(15)
font.bold = True
font.color.rgb = RGBColor(0x00, 0x00, 0x00)

new_heading_style = styles.add_style('New Heading3', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Heading 3']
font = new_heading_style.font
font.name = "Times New Roman"
font.size = Pt(14)
font.bold = True
font.color.rgb = RGBColor(0x00, 0x00, 0x00)

new_heading_style = styles.add_style('Para', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Normal']
font = new_heading_style.font
font.name = "Times New Roman"
font.size = Pt(2)
font.bold = True
font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

sections = doc.sections
doc.add_paragraph('Исследовательский анализ данных (EDA)', style='New Heading1')
tasks = config["tasks"]
i=-1
k=-1
jj=0

for task in tasks:
    i=i+1
    print("folder: ", task["folder"])
    if task['topic'] != " ": 
        header = f"{task['topic']}"
    
        Hstyle = f"{task['Header']}"
        if header != " ": 
            print("Graph Header: ", i, " ", header)
            if Hstyle == "Heading2":
                doc.add_paragraph(header, style='New Heading2') 
            if Hstyle == "Heading3":
                doc.add_paragraph(header, style='New Heading3') 
             
        if (task["Text"] == "y") and (i<len(start)):    
            b=start[i]
            e=stop[i]
            print("start: ", b, " stop: ", e)

            if (e-b) > 2:
                j=0
                for j in range(b, e-1):
                    get_para_data(doc, InDoc.paragraphs[j])
            
        if (task["models"] != " "): 
            models = f"{task['models']}"
            print("models:", models)
            print("path1:", os.path.join(cwd, models))
            path=os.path.join(pathIm, models) 
            print("path2:", path)
                 
            if os.path.isfile(path):
                print("if os.path.isfile(path):", path)
                doc.add_picture(path, width=Inches(5.0))
                  
        if task["Tab"] != " ":     
            jj=jj+1
            tab = f"{task['Tab']}"
            #doc.add_paragraph("Таблица №" + str(jj) + ". Распределение пациентов по заданным категориям", style='New Heading3')  
            #doc.add_paragraph("Таблица №" + str(jj), style='Para')
            doc.add_paragraph("Таблица №" + tab + ". Распределение пациентов по заданным категориям", style='New Heading3')  
            doc.add_paragraph("Таблица №" + tab, style='Para')
        
        if task["sec"] == "y":      
            doc.add_section(WD_SECTION.ODD_PAGE)  

print("ReportEDA: ", ReportEDA) 
doc.save(ReportEDA) 

##########################################
from docx import Document 
document = Document(tabdoc)
doc = Document(ReportEDA) 

count=len(document.tables)
print("Anzahl Tables: ", count)

for num in range (0, count):
    template = document.tables[num]
    print("num: ", num, " template: ", template) 
    tbl = template._tbl
    new_tbl = deepcopy(tbl)
   
    for para in doc.paragraphs:
        if para.text == "Таблица №" + str(num+1): 
            print(para.text)
            para._p.addnext(new_tbl)
            time.sleep(1)

doc.save(ReportAll) 
doc.save(ReportEDAr) 
print(ReportEDAr + str(" Tables fertig!"))
time.sleep(3)

##########################################
document = Document()
doc = Document(ReportEDAr)
jj=0             
for table in doc.tables:
    jj=jj+1
    print("Tables: " + str(jj), table)
    if jj>1:     
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False 
        table.allow_autofit = False
        n_rows=len(table.rows)
        n_cols=len(table.columns)
      
        table.add_row()
        g = table.cell(n_rows, 0)
        h = table.cell(n_rows, n_cols-1)
        g.merge(h)
        cell = table.cell(n_rows, n_cols-1)
        
        cell.paragraphs[0].paragraph_format.space_before = Inches(0)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT 
        cell.paragraphs[0].add_run("© Dr. Alexander Wagner. Все права охраняются законом")
        change_table_cell(table.rows[n_rows].cells[2], background_color="lightgreen", font_color="0000ff", font_size=8, bold=True, italic=True)
    
        for row in table.rows:
            row.height = Cm(0.55)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        table.rows[n_rows-1].height = Cm(0.45)   
        table.rows[n_rows-1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
               
doc.save(ReportEDAr)  
print(ReportEDAr + str(" fertig!"))

######################### Test3.py ##############################
######################### Test3.py, Block 1 #####################
#################################################################
from PIL import Image 

stop=18
df = pd.read_csv(matrix) 
df.rename({'Y': 'target'}, axis=1, inplace=True)
df = df.fillna(0)

train, test = train_test_split(df, test_size = 0.4)

X_train = train[train.columns.difference(['target'])]
y_train = train['target']

X_test = test[test.columns.difference(['target'])]
y_test = test['target']

cv_n_split = 2
random_state = 0
test_train_split_part = 0.2

train0, test0 = X_train, X_test
target0 = y_train
train, test, target, target_test = train_test_split(train0, target0, test_size=test_train_split_part, random_state=random_state)

cv_train = ShuffleSplit(n_splits=cv_n_split, test_size=test_train_split_part, random_state=random_state)
metrics_all = {1 : 'r2_score', 2: 'acc', 3 : 'rmse', 4 : 're'}
metrics_now = [1, 2, 3, 4] 

num_models = 18
acc_train = []
acc_test = []
ShuffleSplit(n_splits=2, random_state=0, test_size=0.2, train_size=None)

############### Alte Modelle ###############
# GradientBoostingClassifier
mGBC=GradientBoostingClassifier()
GBC=GradientBoostingClassifier()
    
# KNeighborsClassifier
mKNC=KNeighborsClassifier()
KNC=KNeighborsClassifier(n_neighbors = 5, metric = 'minkowski', p = 2)

# SVC
mSVC = SVC()
SVC = SVC(kernel="rbf", C=0.025, probability=True)

# DecisionTreeClassifier
mDTC=DecisionTreeClassifier()
DTC=DecisionTreeClassifier(criterion = 'entropy', random_state = 51)

# RandomForestClassifier
mRF=RandomForestClassifier()
RF=RandomForestClassifier(n_estimators = 20, criterion = 'entropy', random_state = 51)

# XGBClassifier
mXGBC=XGBClassifier()
XGBC=XGBClassifier()

# AdaBoostClassifier
mAdaBoost=AdaBoostClassifier()
AdaBoost=AdaBoostClassifier(DecisionTreeClassifier(criterion = 'entropy', random_state = 200),
                              n_estimators=2000,
                              learning_rate=0.1,
                              algorithm='SAMME.R',
                              random_state=1,)    


# LogisticRegression LR
LogReg=LogisticRegression()
LR=LogisticRegression(random_state = 51)

############### Neue Modelle ###############
# Linear Regression
param_grid = {'C': np.linspace(.1, 1.5, 15)}
linreg = LogisticRegression()
linreg_CV = GridSearchCV(linreg, param_grid=param_grid, cv=cv_train, verbose=False)

# Linear SVR
linear_svc = LinearSVC()
linear_svc_CV = GridSearchCV(linear_svc, param_grid={}, cv=cv_train, verbose=False)

mlp = MLPClassifier()
param_grid = {'hidden_layer_sizes': [i for i in range(2,5)],
              'solver': ['sgd'],
              'learning_rate': ['adaptive'],
              'max_iter': [1000]
              }
mlp_GS = GridSearchCV(mlp, param_grid=param_grid, cv=cv_train, verbose=False)

# Decision Tree Classifier
decision_tree = DecisionTreeClassifier()
decision_tree_CV = GridSearchCV(decision_tree, param_grid={}, cv=cv_train, verbose=False)

# Stochastic Gradient Descent
sgd = SGDClassifier()
sgd_CV = GridSearchCV(sgd, param_grid={}, cv=cv_train, verbose=False)

# Ridge Classifier
ridge = RidgeClassifier()
ridge_CV = GridSearchCV(estimator=ridge, param_grid={'alpha': np.linspace(.1, 1.5, 15)}, cv=cv_train, verbose=False)

# Bagging Classifier
bagging = BaggingClassifier()
bagging_CV = GridSearchCV(estimator=bagging, param_grid={}, cv=cv_train, verbose=False)

# AdaBoost Classifier
Ada_Boost = AdaBoostClassifier()
Ada_Boost_CV = GridSearchCV(estimator=Ada_Boost, param_grid={'learning_rate' : [.01, .1, .5, 1]}, cv=cv_train, verbose=False)

# LogisticRegression
logreg = LogisticRegression()
logreg_CV = GridSearchCV(estimator=logreg, param_grid={'C' : [.1, .3, .5, .7, 1]}, cv=cv_train, verbose=False)

# Gaussian Naive Bayes
gaussian = GaussianNB()
gaussian_CV = GridSearchCV(estimator=gaussian, param_grid={}, cv=cv_train, verbose=False)

# Perceptron
perceptron = Perceptron()
perceptron_CV = GridSearchCV(estimator=perceptron, param_grid={}, cv=cv_train, verbose=False)

classifiers = [
    linreg_CV,
    logreg_CV,
    perceptron_CV,
    linear_svc_CV,
    mlp_GS,
    decision_tree_CV,
    sgd_CV,
    ridge_CV,
    bagging_CV,
    Ada_Boost_CV,
    GBC,
    KNC,
    DTC,
    RF,
    XGBC,
    AdaBoost,
    gaussian_CV,
    SVC
  ] 

models = pd.DataFrame({'Model': [
    'Linear Regression', 
    'Logistic Regression',
    'Perceptron',
    'Linear SVC', 
    'MLPClassifier', 
    'Decision Tree Classifier 1', 
    'Stochastic Gradient Decent', 
    'RidgeClassifier', 
    'BaggingClassifier', 
    'AdaBoostClassifier 1', 
    'GradientBoostingClassifie',
    'KNeighborsClassifier',
    'DecisionTreeClassifier 2',
    'RandomForestClassifier',
    'XGBClassifier',
    'AdaBoostClassifier 2',
    'Naive Bayes',    
    'SVC' ]})

i=-1
result_table = pd.DataFrame(columns=['classifiers', 
                                     'fpr',
                                     'tpr',
                                     'accurate',
                                     'auc'])
for classifier in classifiers:
    i=i+1
    name_model = models.iloc[i]['Model']
    if i < stop:
        pipe = Pipeline(steps=[('classifier', classifier)])
        model = pipe.fit(X_train, y_train) 
        acc_metrics_calc(i,model,train,test,target,target_test)

        title = "Model: " + name_model
        print("i: ", i, title)
        figure, axes = plt.subplots(2, 1, figsize=(20, 10))
    
        if axes is None:
            _, axes = plt.subplots(1, 2, figsize=(20, 5))

        axes[0].set_title(title)
        axes[0].set_xlabel("Training examples")
        axes[0].set_ylabel("Score")

        cv_train = ShuffleSplit(n_splits=cv_n_split, test_size=test_train_split_part, random_state=random_state)
        train_sizes, train_scores, test_scores, fit_times, _ = \
            learning_curve(estimator=model, X=train, y=target, cv=cv_train,
                           train_sizes=np.linspace(.1, 1.0, 5),
                           return_times=True)
        train_scores_mean = np.mean(train_scores, axis=1)
        train_scores_std = np.std(train_scores, axis=1)
        test_scores_mean = np.mean(test_scores, axis=1)
        test_scores_std = np.std(test_scores, axis=1)
        fit_times_mean = np.mean(fit_times, axis=1)
        fit_times_std = np.std(fit_times, axis=1)

        axes[0].grid()
        axes[0].fill_between(train_sizes, train_scores_mean - train_scores_std,
                             train_scores_mean + train_scores_std, alpha=0.1,
                             color="r")
        axes[0].fill_between(train_sizes, test_scores_mean - test_scores_std,
                             test_scores_mean + test_scores_std, alpha=0.1,
                             color="g")
        axes[0].plot(train_sizes, train_scores_mean, 'o-', color="r",
                     label="Training score")
        axes[0].plot(train_sizes, test_scores_mean, 'o-', color="g",
                     label="Cross-validation score")
        axes[0].legend(loc="best")

        # Plot n_samples vs fit_times
        axes[1].grid()
        axes[1].plot(train_sizes, fit_times_mean, 'o-')
        axes[1].fill_between(train_sizes, fit_times_mean - fit_times_std,
                             fit_times_mean + fit_times_std, alpha=0.1)
        axes[1].set_xlabel("Training examples")
        axes[1].set_ylabel("fit_times")
        axes[1].set_title("Scalability of the model")

        figure.savefig(pathIm + '/plot_learning_curve' + str(i) + '.png')
        ypred   = model.predict(X_test)
         
        if (i==2): 
            clf = CalibratedClassifierCV(perceptron_CV) 
            clf.fit(X_train, y_train)
            y_pred = clf.predict_proba(X_test)[::,1]
        elif (i==3):
            #svm = LinearSVC()
            clf = CalibratedClassifierCV(linear_svc_CV) #svm) 
            clf.fit(X_train, y_train)
            y_pred = clf.predict_proba(X_test)[::,1]
        elif (i==6):
            clf = CalibratedClassifierCV(sgd_CV)
            clf.fit(X_train, y_train)
            y_pred = clf.predict_proba(X_test)[::,1]            
        elif (i==7):
            clf = CalibratedClassifierCV(ridge_CV)
            clf.fit(X_train, y_train)
            y_pred = clf.predict_proba(X_test)[::,1]            
               
        else:
            y_pred = model.predict_proba(X_test)[::,1]
  
            
        fpr, tpr, _ = roc_curve(y_test, y_pred) 
        auc         = roc_auc_score(y_test, y_pred).round(3) 
        accuracy    = accuracy_score(ypred, y_test).round(2)
        F1_Score    = f1_score(y_test, ypred).round(2)
        Precision_Score = precision_score(y_test, ypred).round(2)
        Recall_Score = recall_score(y_test, ypred).round(2)
        Balanced_Accuracy_Score = balanced_accuracy_score(y_test, ypred).round(2)
        
        target_names = ['class 0', 'class 1']
        report = classification_report(y_test, ypred, target_names=target_names, output_dict=True, digits=4)
        print("report:", report)
        CR = pd.DataFrame(report).transpose()
        CR['AUC'] = auc 
        print("CR: ", CR)
        CR.to_csv(pfad + '/CLSB_{i}.csv')
                    
        Accurate_train=pipe.score(X_train, y_train).round(2)   
        Accurate_test =pipe.score(X_test,  y_test).round(2)  
        cm = confusion_matrix(y_test, ypred)
        fig=plt.figure(figsize=(8,5))
        plt.title('Heatmap of Confusion Matrix', fontsize = 15)
        sns.heatmap(cm, annot = True)
        fig.savefig(pathIm + '/Heatmap' + str(i) + '.png') 
       
        
        Roc_Auc_Score = roc_auc_score(y_test, y_pred).round(2)
        fig2=plt.figure(figsize=(8,5))            
        plt.style.use('seaborn-v0_8-darkgrid')
        plt.plot(fpr, tpr, label='%s ROC (area = %0.2f)' % (name_model, auc))
        plt.plot([0, 1], [0, 1],'r--')
        plt.xlim([0.0, 1.0])
        plt.ylim([0.0, 1.05])
        plt.xlabel('1-Specificity(False Positive Rate)')
        plt.ylabel('Sensitivity(True Positive Rate)')

        plt.legend(loc="lower right")
        #plt.show() 
        fig2.savefig(pathIm + '/PlotROC' + str(i) + '.png') 

i=-1
plt.style.use('seaborn-v0_8-darkgrid')
figC = plt.figure(figsize=(16,10))

for classifier in classifiers:
    i=i+1
    name_model = models.iloc[i]['Model']
  
    if i < stop:
        pipe = Pipeline(steps=[('classifier', classifier)])
        model = pipe.fit(X_train, y_train) 
        ypred   = model.predict(X_test)
        if (i==2): 
            clf = CalibratedClassifierCV(perceptron_CV) 
            clf.fit(X_train, y_train)
            y_pred = clf.predict_proba(X_test)[::,1]
        elif (i==3):
            #svm = LinearSVC()
            clf = CalibratedClassifierCV(linear_svc_CV) #svm) 
            clf.fit(X_train, y_train)
            y_pred = clf.predict_proba(X_test)[::,1]
        elif (i==6):
            clf = CalibratedClassifierCV(sgd_CV)
            clf.fit(X_train, y_train)
            y_pred = clf.predict_proba(X_test)[::,1]            
        elif (i==7):
            clf = CalibratedClassifierCV(ridge_CV)
            clf.fit(X_train, y_train)
            y_pred = clf.predict_proba(X_test)[::,1]            
               
        else:
            y_pred = model.predict_proba(X_test)[::,1]
             
        fpr, tpr, _ = roc_curve(y_test, y_pred) 
        auc         = roc_auc_score(y_test, y_pred).round(3) 
        accuracy    = accuracy_score(ypred, y_test).round(2)
        F1_Score    = f1_score(y_test, ypred).round(2)

        Precision_Score = precision_score(y_test, ypred).round(2)
        Recall_Score = recall_score(y_test, ypred).round(2)
        Balanced_Accuracy_Score = balanced_accuracy_score(y_test, ypred).round(2)
        Classification_Report = classification_report(y_test, ypred)
        Accurate_train=pipe.score(X_train, y_train).round(2)   
        Accurate_test =pipe.score(X_test,  y_test).round(2)        
         
        plt.plot(fpr, tpr, label='%s ROC (area = %0.2f)' % (name_model, auc))
        plt.plot([0, 1], [0, 1],'r--')
        plt.xlim([0.0, 1.0])
        plt.ylim([0.0, 1.05])
        plt.xlabel('1-Specificity(False Positive Rate)')
        plt.ylabel('Sensitivity(True Positive Rate)')
        plt.title('Receiver Operating Characteristic')
        plt.legend(loc="lower right")
        
img = fig2img(figC) 
img.save(pathIm + '/PlotRoc.png') 

#6. Speicher der Ergebnisse
metricsnow = pd.DataFrame(metrics_now)
metricsnow.to_csv(pfad + '/Nostalgi2023_metrics_now.csv', index=False) 
accall= pd.DataFrame(acc_all)
print("acc_all:", acc_all)
accall.to_csv(pfad + '/RT4.csv', index=False)

models = pd.DataFrame({'Model': [
    'Linear Regression', 
    'Logistic Regression',
    'Perceptron',
    'Linear SVC', 
    'MLPClassifier', 
    'Decision Tree Classifier 1', 
    'Stochastic Gradient Decent', 
    'RidgeClassifier', 
    'BaggingClassifier', 
    'AdaBoostClassifier 1', 
    'GradientBoostingClassifie',
    'KNeighborsClassifier',
    'DecisionTreeClassifier 2',
    'RandomForestClassifier',
    'XGBClassifier',
    'AdaBoostClassifier 2',
    'Naive Bayes',    
    'SVC' ]})

metrics_all = {1 : 'r2_score', 2: 'acc', 3 : 'rmse', 4 : 're'}
metrics_now = [1, 2, 3, 4]
models=models[:] 
models['Nr'] = models.reset_index().index
print(models)

for x in metrics_now:
    xs = metrics_all[x]
    models[xs + '_Train'] = acc_all[(x-1)*2]
    models[xs + '_Test']  = acc_all[(x-1)*2+1]
    if xs == "acc":
        models[xs + '_Diff'] = models[xs + '_Train'] - models[xs + '_Test']

ms = metrics_all[metrics_now[1]] # the accuracy
models.sort_values(by=[(ms + '_Test'), (ms + '_Train')], ascending=False)
pd.options.display.float_format = '{:,.2f}'.format

# Plots
plt.style.use('seaborn-v0_8-darkgrid')

i=0
for x in metrics_now: 
    i=i+1
    xs = metrics_all[x]
    xs_train = metrics_all[x] + '_Train'
    xs_test = metrics_all[x] + '_Test'
    figD = plt.figure(figsize=[25,6])
    xx = models['Model']
    plt.tick_params(labelsize=14)
    plt.plot(xx, models[xs_train], label = xs_train)
    plt.plot(xx, models[xs_test], label = xs_test)
    plt.legend()
    plt.title(str(xs) + ' criterion for ' + str(num_models) + ' popular models for train and test datasets')
    plt.xlabel('Models')
    plt.ylabel(xs + ', %')
    plt.xticks(xx, rotation='vertical')
    img = fig2img(figD) 
    img.save(pathIm + "/PlotE" + str(i) + ".png")

metrics_main = 2
xs = metrics_all[metrics_main]
xs_train = metrics_all[metrics_main] + '_Train'
xs_test = metrics_all[metrics_main] + '_Test'
direct_sort = False if (metrics_main >= 2) else True

models_sort = models.sort_values(by=[xs_test, xs_train], ascending=direct_sort)
models_best = models_sort[(models_sort.acc_Diff < 5) & (models_sort.acc_Train > 90)]
models_best[['Model', ms + '_Train', ms + '_Test', 'acc_Diff']].sort_values(by=['acc_Test'], ascending=False)
print("models_best RT1: eins!", models_best)
models_best.to_csv(pfad + '/RT1.csv', index=False) 

metrics_all = {1 : 'r2_score', 2: 'acc', 3 : 'rmse', 4 : 're'}
metrics_now = [1, 2, 3, 4] 

# list of accuracy of all model - amount of metrics_now * 2 (train & test datasets)
num_models = 18
acc_train = []
acc_test = []
acc_all = np.empty((len(metrics_now)*2, 0)).tolist()

models = pd.DataFrame({'Model': [
    'Linear Regression', 
    'Logistic Regression',
    'Perceptron',
    'Linear SVC', 
    'MLPClassifier', 
    'Decision Tree Classifier 1', 
    'Stochastic Gradient Decent', 
    'RidgeClassifier', 
    'BaggingClassifier', 
    'AdaBoostClassifier 1', 
    'GradientBoostingClassifie',
    'KNeighborsClassifier',
    'DecisionTreeClassifier 2',
    'RandomForestClassifier',
    'XGBClassifier',
    'AdaBoostClassifier 2',
    'SVC'
 ]})

models=models[:] 
models['Nr'] = models.reset_index().index
print(models)

models_pred = pd.DataFrame(models.Model, columns = ['Model'])
N_best_models = len(models_pred.Model)

from sklearn.svm import SVC
i=0
for i in range(N_best_models):
    if i < N_best_models+1:
        name_model = models_pred.iloc[i]['Model']
        ii=i+1
        model = model_fit(name_model,train0,target0)
        acc_metrics_calc_pred(i,model,name_model, train0, test0, target0) #Def in AsniDef.py

models_pred=models_pred[0:N_best_models]

i=0
for x in metrics_now:
    xs = metrics_all[x]
    #Auswahl N_best_models(17 im Moment) Spalten aus der acc_all_pred!  
    acc_all_pred2=acc_all_pred[x-1][0:N_best_models] 
    models_pred[xs + '_train'] = acc_all_pred2 

sort_pred=models_pred[['Model', 'acc_train']].sort_values(by=['acc_train'], ascending=False)
sort_pred.to_csv(pfad + '/RT2.csv', index=False) 

models_pred=models_pred[0:N_best_models]
i=0
for x in metrics_now:
    xs = metrics_all[x]
    #Auswahl N_best_models(17 im Moment) Spalten aus der acc_all_pred!  
    acc_all_pred2=acc_all_pred[x-1][0:N_best_models] 
    models_pred[xs + '_train'] = acc_all_pred2 

print(models_pred)
sort_pred=models_pred[['Model', 'r2_score_train','acc_train','rmse_train', 're_train']].sort_values(by=['acc_train'], ascending=False)
sort_pred.to_csv(pfad + '/RT3.csv', index=False) 

#import pandas as pd
RT1=pd.read_csv(pfad + "\RT1.csv")
RT2=pd.read_csv(pfad + "\RT2.csv")
RT3=pd.read_csv(pfad + "\RT3.csv")
RT4=pd.read_csv(pfad + "\RT4.csv")
RT5=pd.read_csv(pfad + "\pred2023.csv")
RT5.to_csv(pfad + '/RT5.csv', index=False) 

print('models_best RT1', RT1) #models_best)
print('------------------------------------------------------------------------- ')
print(' ')
print('sort_pred RT2', RT2) #sort_pred)
print('------------------------------------------------------------------------- ')
print(' ')
print('models_pred RT3', RT3) #models_pred)
print('------------------------------------------------------------------------- ')
print('models_best RT4', RT4) #accall
print('------------------------------------------------------------------------- ')
print('models_best RT5', RT5)
print('------------------------------------------------------------------------- ')

######################### Test3.py, Block 2 #####################
#################################################################
list = [5, 8, 10, 12, 13, 14, 15]
metrics_all = {1 : 'precision', 2: 'recall', 3 : 'f1-score', 4 : 'AUC'}
metrics_now = [1, 2, 3, 4] 

#global acc_all
acc_all = np.empty((len(metrics_now), 0)).tolist()

for i in list:
    data=pd.read_csv(pfad + '/CLSB_' + str(i) + ".csv")
    headers =  ["metriks","precision", "recall", "f1-score", "support", "AUC"]
    data.columns = headers
    print(data)
    num_acc = 0
    for x in metrics_now:
        if x == 1:
           acc_train = data.loc[data['metriks'] == 'weighted avg']['precision'].values[0]
        elif x == 2:
            acc_train = data.loc[data['metriks'] == 'weighted avg']['recall'].values[0]
        elif x == 3:
           acc_train =  data.loc[data['metriks'] == 'weighted avg']['f1-score'].values[0]
        elif x == 4:
           acc_train =  data.loc[data['metriks'] == 'weighted avg']['AUC'].values[0]   
        acc_all[num_acc].append(acc_train) 
        num_acc += 1
    
accall= pd.DataFrame(acc_all)
headers =  ["Var1", "Var2", "Var3", "Var4" , "Var5" , "Var6", "Var7"] 
accall.columns = headers
l_bool = [True, False, False, False]
print(accall[l_bool])
l_bool = [False, True, False, False]
print(accall[l_bool])
l_bool = [False, False, False, True]
da=accall.loc[[0], :]
print("3", da)

Rfile = pfad + '/Ametriks.csv'
accall.to_csv(Rfile, index=False)   

######################### Test3.py, Block 3 #####################
#################################################################
''' 26.08.2024
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from PIL import Image 

out=Dstr() ?????????????????
col1=out[1]
col2=out[2]
dfm=Combined(col2, col1)
graphM(dfm, pathIm)
'''

######################### Test4.py ##############################
#################################################################
### 1978

today = datetime.date.today()
year = today.year
cwd = os.getcwd()
print("Start Step4!")
template=DirTemplate
CONFIG_JSON = config_json #os.path.join(cwd, "Templates\ASNIR.json")
ASSETS = DirAssets #os.path.join(DirPr, "ASSETS")

reportWordPath = os.path.join(DirPr, 'OUTPUT\ASNI_ReportTest01.docx')
TemplateTOC = os.path.join(DirTemplate, 'TemplateTOC.docx')

print("template: ", template)
print("rReportTest: ", rReportTest)
print("reportWordPath: ", reportWordPath)
print("TemplateTOC: ", TemplateTOC)

from docxtpl import DocxTemplate
Inmodels = pd.DataFrame({'Model': [
   'Linear Regression', 
   'Logistic Regression',
   'Perceptron',
   'Linear SVC', 
   'MLPClassifier', 
   'Decision Tree Classifier 1', 
   'Stochastic Gradient Decent', 
   'RidgeClassifier', 
   'BaggingClassifier', 
   'AdaBoostClassifier 1', 
   'GradientBoostingClassifie',
   'KNeighborsClassifier',
   'DecisionTreeClassifier 2',
   'RandomForestClassifier',
   'XGBClassifier',
   'AdaBoostClassifier 2',
   'Naive Bayes',    
   'SVC' ]})    

Nk=len(Inmodels)

#Spire
from spire.doc.common import *
from spire.doc import *

# Используемые стили
HEADER_STYLE = "BoldHeader"
HEADER_LINK_STYLE = "BoldHeaderHyperlink"
CONTENT_STYLE = "Content"
CODE_STYLE = "Code"

document = Document()
section = document.AddSection()
paragraph = section.AddParagraph()
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text =paragraph.AppendText("{{Projekt}}")  
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 26
text.CharacterFormat.Bold = True
text.CharacterFormat.TextColor = Color.get_Blue()

paragraph = section.AddParagraph()
paragraph = section.AddParagraph() 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("Тема исследования: {{Projekt3}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 16
text.CharacterFormat.Bold = True

paragraph = section.AddParagraph()
paragraph = section.AddParagraph() 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("Проект: {{Thema}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 16
text.CharacterFormat.Bold = True

paragraph = section.AddParagraph()
paragraph = section.AddParagraph()  
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("Автор исследования: {{Forscher}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 16
text.CharacterFormat.Bold = True
for num in range(9):
    i=num+1
    paragraph = section.AddParagraph()  
    
paragraph = section.AddParagraph()
paragraph = section.AddParagraph() 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("{{logo}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 16
text.CharacterFormat.Bold = True

for num in range(13):
    i=num+1
    paragraph = section.AddParagraph()  
    
paragraph = section.AddParagraph()  
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("{{Site}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 14
text.CharacterFormat.Bold = True

paragraph = section.AddParagraph() 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("{{Year}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 14
text.CharacterFormat.Bold = True

section = document.AddSection()
paragraph = section.AddParagraph()
paragraph = section.AddParagraph()

paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("Содержание") 
text.CharacterFormat.FontName = "Arial"
text.CharacterFormat.FontSize = 14
text.CharacterFormat.Bold = True
text.CharacterFormat.Italic = True

paragraph = section.AddParagraph()
text = paragraph.AppendText(" ") 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center

#paragraph = section.AddParagraph()
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
outputFile = template + "/TemplateTOC1.docx"  
document.SaveToFile(outputFile, FileFormat.Docx)
document.Close()

from docx import Document, enum
doc = Document(TemplateTOC)
lines = doc.paragraphs
for line in lines:
    if "{{ClassBlk" in line.text or "{{Kap" in line.text: 
        print(line.text)   
        line.paragraph_format.first_line_indent = Inches(0.25)
        continue

time.sleep(2.4)
doc.save(rReportTest)

#################################################################
#################################################################
### 2175
#cwd = os.getcwd()
#template=os.path.join(cwd, "Templates")

'''
def replace_copy(txt):
    finder = wordapp.Selection.Find
    finder.Text = txt 
    finder.Execute()
    wordapp.Selection.MoveStart
    wordapp.Selection.Paste()
    if txt=="KapEDA":
        time.sleep(10.0)
'''

today = datetime.date.today()
year = today.year

print("Start Step4!")
from docxtpl import DocxTemplate
Inmodels = pd.DataFrame({'Model': [
   'Linear Regression', 
   'Logistic Regression',
   'Perceptron',
   'Linear SVC', 
   'MLPClassifier', 
   'Decision Tree Classifier 1', 
   'Stochastic Gradient Decent', 
   'RidgeClassifier', 
   'BaggingClassifier', 
   'AdaBoostClassifier 1', 
   'GradientBoostingClassifie',
   'KNeighborsClassifier',
   'DecisionTreeClassifier 2',
   'RandomForestClassifier',
   'XGBClassifier',
   'AdaBoostClassifier 2',
   'Naive Bayes',    
   'SVC' ]})    

Nk=len(Inmodels)
#Spire
from spire.doc.common import *
from spire.doc import *

# Используемые стили
HEADER_STYLE = "BoldHeader"
HEADER_LINK_STYLE = "BoldHeaderHyperlink"
CONTENT_STYLE = "Content"
CODE_STYLE = "Code"

document = Document()
section = document.AddSection()
paragraph = section.AddParagraph()
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text =paragraph.AppendText("{{Projekt}}")  
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 26
text.CharacterFormat.Bold = True
text.CharacterFormat.TextColor = Color.get_Blue()

paragraph = section.AddParagraph()
paragraph = section.AddParagraph() 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("Тема исследования: {{Projekt3}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 16
text.CharacterFormat.Bold = True

paragraph = section.AddParagraph()
paragraph = section.AddParagraph() 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("Проект: {{Thema}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 16
text.CharacterFormat.Bold = True

paragraph = section.AddParagraph()
paragraph = section.AddParagraph()  
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("Автор исследования: {{Forscher}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 16
text.CharacterFormat.Bold = True
for num in range(9):
    i=num+1
    paragraph = section.AddParagraph()  
    
paragraph = section.AddParagraph()
paragraph = section.AddParagraph() 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("{{logo}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 16
text.CharacterFormat.Bold = True

for num in range(13):
    i=num+1
    paragraph = section.AddParagraph()  
    
paragraph = section.AddParagraph()  
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("{{Site}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 14
text.CharacterFormat.Bold = True

paragraph = section.AddParagraph() 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("{{Year}}")
text.CharacterFormat.FontName = "Times New Roman"
text.CharacterFormat.FontSize = 14
text.CharacterFormat.Bold = True

section = document.AddSection()
paragraph = section.AddParagraph()
paragraph = section.AddParagraph()

paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center
text = paragraph.AppendText("Содержание") 
text.CharacterFormat.FontName = "Arial"
text.CharacterFormat.FontSize = 14
text.CharacterFormat.Bold = True
text.CharacterFormat.Italic = True

paragraph = section.AddParagraph()
text = paragraph.AppendText(" ") 
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center

#paragraph = section.AddParagraph()
paragraph.Format.HorizontalAlignment = HorizontalAlignment.Center

#template=os.path.join(cwd, "Templates")
outputFile = template + "/TemplateTOC1.docx"  
document.SaveToFile(outputFile, FileFormat.Docx)
document.Close()

from docx import Document, enum
doc = Document(TemplateTOC)
lines = doc.paragraphs
for line in lines:
    if "{{ClassBlk" in line.text or "{{Kap" in line.text: 
        print(line.text)   
        line.paragraph_format.first_line_indent = Inches(0.25)
        continue

time.sleep(2.4)
doc.save(rReportTest) 

doc = DocxTemplate(rReportTest)
Nk=18

with open(os.path.join(DirTemplate, 'ASNIR.json'), 'r', encoding='utf-8') as file_object:    
    ASNI_dict = json.load(file_object)

ASNI_dict['logo'] = InlineImage(doc,  os.path.join(DirAssets, 'ImageAll\GRAPH0.png'), Cm(12))  

for num in range(Nk):
    i=num+1
    ASNI_dict['Heatmap'+str(num)] = InlineImage(doc, os.path.join(DirAssets, 'ImageAll\Heatmap' + str(num) + '.png'), Cm(18)) 
    ASNI_dict['PltR' + str(num)]  = InlineImage(doc, os.path.join(DirAssets, 'ImageAll\PlotROC' + str(num) + '.png'), Cm(18))  
    ASNI_dict['PltL' + str(num)]  = InlineImage(doc, os.path.join(DirAssets, 'ImageAll\plot_learning_curve' + str(num+1) + '.png'), Cm(20))  

doc.render(ASNI_dict)
doc.save(reportWordPath) 
ReportTest01= os.path.join(DirPr, "OUTPUT\ASNI_ReportTest01.docx")
print("ReportTest01: ", ReportTest01)

word = win32com.client.gencache.EnsureDispatch("Word.Application")
word.Visible = False  
doc = word.Documents.Open(ReportTest01)  
i=-1

for num in range(Nk):
    i=i+1
    print('CLSB: ', i, 'CLSB_'+str(i)+'.csv')
    df=pd.read_csv(os.path.join(pfad, 'CLSB_'+str(i)+'.csv'))
    rng = doc.Bookmarks("Table" +  str(num)).Range
    Table=rng.Tables.Add(rng,NumRows=df.shape[0]+1,NumColumns=df.shape[1])
    
    for col in range(df.shape[1]):        
        Table.Cell(1,col+1).Range.Text=str(df.columns[col]) 
        for row in range(df.shape[0]):
            Table.Cell(row+1+1,col+1).Range.Text=str(df.iloc[row,col])  
doc.Close()
word.Quit()
print("Table in Ordnung!")

from docx import Document, enum
document = Document()
doc = Document(os.path.join(DirPr, "OUTPUT\ASNI_ReportTest01.docx"))

for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False 
    table.allow_autofit = False
    n_rows=len(table.rows)
    n_cols=len(table.columns)
        
    table.cell(0, 0).text = 'Classes+Metrics'
    table.add_row()
    set_column_width(table, 0, 35)
    for c in range(1, 5):
        set_column_width(table, c, 20)

    
    g = table.cell(n_rows, 0)
    h = table.cell(n_rows, n_cols-1)
    g.merge(h)
    cell = table.cell(n_rows, n_cols-1)
    
    cell.paragraphs[0].paragraph_format.space_before = Inches(0)
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT 
    cell.paragraphs[0].add_run("© Dr. Alexander Wagner. Все права охраняются законом")
    change_table_cell(table.rows[n_rows].cells[2], background_color="lightgreen", font_color="0000ff", font_size=8, bold=True, italic=True)
    table.style = 'Table Grid' 
    set_repeat_table_header(table.rows[0])

    for i in range(1, n_rows):
        for j in range(1, n_cols):
            element=table.cell(i, j).text
            partition = element.partition('.')
            if (partition[0].isdigit() and partition[1] == '.' and partition[2].isdigit()):
                newelement =  float(element) 
                y=round(newelement,3)
                table.cell(i, j).text=str(y)
                table.cell(i, j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
    
    for c in range(0, n_cols):
        change_table_cell(table.rows[0].cells[c], background_color="lightgreen", font_color="0000ff", font_size=12, bold=True, italic=True)             
        
        for cell in table.columns[c].cells:  
            cell.paragraphs[0].paragraph_format.space_after = Inches(0)
            cell.paragraphs[0].paragraph_format.space_before = Inches(0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
       
            set_cell_border(
                cell,
                top={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
                bottom={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
                left={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
                right={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
                insideH={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
                end={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"}
            )
       
     
    for row in table.rows:
        row.height = Cm(0.55)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
    table.rows[0].height = Cm(0.6)   
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    table.rows[n_rows-1].height = Cm(0.45)   
    table.rows[n_rows-1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
 
time.sleep(2.4)   
doc.save(os.path.join(DirPr, "ASNI_Report.docx"))  
print("ASNI_Report.docx fertig!")

from spire.doc import *
document = Document()
file = os.path.join(DirPr, "ASNI_Report.docx")
dt=datetime.datetime.fromtimestamp(os.stat(file).st_mtime)

txd = "Документ актуализирован: " + dt.strftime('%d.%m.%Y %H:%M:%S')
print("Mody:", txd)

# Load a Word document 
document.LoadFromFile(file)
# Get the first section
section = document.Sections[0]

# Get header
header = section.HeadersFooters.Header

# Add a paragraph to the header and set its alignment style
headerParagraph = header.AddParagraph() 
headerParagraph.Format.HorizontalAlignment = HorizontalAlignment.Left
#headerParagraph.Format.VerticalAlignment = VerticalAlignment.Center
section.header_distance = Cm(1.2)

headerPicture = headerParagraph.AppendPicture(os.path.join(ASSETS, "ImageAll/logo2.jpg"))
headerPicture.TextWrappingStyle = TextWrappingStyle.Square
headerPicture.VerticalOrigin = VerticalOrigin.Line
headerPicture.VerticalAlignment = ShapeVerticalAlignment.Center
#headerPicture.HorizontalAlignment = ShapeHorizontalAlignment.Right
headerPicture.HorizontalAlignment = ShapeHorizontalAlignment.Left
headerPicture.VerticalOrigin = VerticalOrigin.TopMarginArea

text = headerParagraph.AppendText("Автоматизировання Система Научных Исследований в медицине и здравоохранении «АСНИ-МЕД»")
text.CharacterFormat.FontName = "Times New"
text.CharacterFormat.FontSize = 9
text.CharacterFormat.Bold = True
text.CharacterFormat.TextColor = Color.get_Blue()

section = document.Sections[0]
# Get footer
footer = section.HeadersFooters.Footer

# Add a paragraph to the footer paragraph and set its alignment style
footerParagraph = footer.AddParagraph()
footerParagraph.Format.HorizontalAlignment = HorizontalAlignment.Left 
# Add text to the footer paragraph and set its font style
text = footerParagraph.AppendText("© Dr. Alexander Wagner, Все права охраняются законом. " + txd)
text.CharacterFormat.FontName = "Times New"
text.CharacterFormat.FontSize = 9
text.CharacterFormat.Bold = True
text.CharacterFormat.TextColor = Color.get_Blue()

footerParagraph = footer.AddParagraph()
footerParagraph.Format.HorizontalAlignment = HorizontalAlignment.Right 
text = footerParagraph.AppendText("Page ")
txt1=footerParagraph.AppendField("page number", FieldType.FieldPage)
txt2=footerParagraph.AppendText(" of ")
txt3=footerParagraph.AppendField("number of pages", FieldType.FieldNumPages)
text.CharacterFormat.TextColor = Color.get_Blue()
txt1.CharacterFormat.TextColor = Color.get_Blue()
txt2.CharacterFormat.TextColor = Color.get_Blue()
txt3.CharacterFormat.TextColor = Color.get_Blue()

# Save the result file
time.sleep(2.4)
document.SaveToFile(os.path.join(DirPr, "AddFootnoteForParagraph.docx"), FileFormat.Docx2016)
document.Close()
time.sleep(2.4)

#def delete_paragraph(paragraph):
from docx import Document
doc = Document(os.path.join(DirPr, "AddFootnoteForParagraph.docx"))
s=len(doc.sections) 
for nt in range(s):    
    section = doc.sections[nt]
    header = doc.sections[nt].header
    footer = doc.sections[nt].footer
    
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    
    header_para  = header.paragraphs[0]
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after  = Pt(7)
    
    footer_para  = footer.paragraphs[0]
    footer_para.paragraph_format.space_before = Pt(0)
    footer_para.paragraph_format.space_after  = Pt(0)
    
    footer_para  = footer.paragraphs[1]
    footer_para.paragraph_format.space_before = Pt(0)
    footer_para.paragraph_format.space_after  = Pt(0)

section = doc.sections[0]  
section.different_first_page_header_footer = True

lines = doc.paragraphs
n=-1
for line in lines:
    n=n+1
    if line.text == "Evaluation Warning: The document was created with Spire.Doc for Python.":    
        delete_paragraph(line)
        continue
time.sleep(3.4)

reportWordPath = os.path.join(DirPr, "ASNI_ReportPre.docx")
doc.save(os.path.join(DirPr, "ASNI_ReportPre.docx"))
print("ASNI_ReportPre.docx fertig!")
time.sleep(2.4) 

######################### Test4RR.py ############################
#################################################################
### 2591

today = datetime.date.today()
year = today.year
Inmodels = pd.DataFrame({'Model': [
   'Linear Regression', 
   'Logistic Regression',
   'Perceptron',
   'Linear SVC', 
   'MLPClassifier', 
   'Decision Tree Classifier 1', 
   'Stochastic Gradient Decent', 
   'RidgeClassifier', 
   'BaggingClassifier', 
   'AdaBoostClassifier 1', 
   'GradientBoostingClassifie',
   'KNeighborsClassifier',
   'DecisionTreeClassifier 2',
   'RandomForestClassifier',
   'XGBClassifier',
   'AdaBoostClassifier 2',
   'Naive Bayes',    
   'SVC' ]})    
Nk=len(Inmodels)
print(Nk)

#####################################
from win32com.client import constants
import win32com.client, time, pythoncom
from win32com import client

'''
def replace_copy(txt):
    finder = wordapp.Selection.Find
    finder.Text = txt 
    finder.Execute()
    #wordapp.Selection.MoveLeft()
    #wordapp.Selection.MoveDown()
    wordapp.Selection.MoveStart
    wordapp.Selection.Paste()
    if txt=="KapEDA":
        time.sleep(10.0)
'''

print("ReportPre: ", reportWordPath)
KapE = os.path.join(pfad, "Kap")
print("KapE: ", KapE)
KapM = os.path.join(pfad, "Modr")
print("KapM: ", KapM)
ResultDoc = os.path.join(DirPr, "OUTPUT\ASNI_Result2025.docx")
print("ResultDoc: ", ResultDoc)
time.sleep(2.4)

KapM = os.path.join(pfad, "Modr")
print("KapM: ", KapM)

wordapp = win32com.client.gencache.EnsureDispatch("Word.Application")
wordapp.Visible = False #True 
wordapp.DisplayAlerts = 0
newdoc = wordapp.Documents.Open(reportWordPath) 
for k in range(0, 14):    
    Text_copy(KapE + str(k) + ".docx") 
    time.sleep(4.4)
    KapTxt="Kap" + str(k)
    print(KapTxt)
    replace_copy(KapTxt) 

newdoc.Select()
finder = wordapp.Selection.Find
finder.Text = "Предисловие"
finder.Execute()
wordapp.Selection.MoveLeft()

for num in range(Nk):
    print(KapM + str(num+1) + ".docx")
    Text_copy(KapM + str(num+1) + ".docx") 
    time.sleep(4.4)
    MoTxt="ClassBlk" + str(num)
    print(num, " ", MoTxt)
    replace_copy(MoTxt)

time.sleep(44.4)
update_tocHalb(newdoc)
newdoc.SaveAs(ResultDoc)
wordapp.Application.Quit(-1)    

time.sleep(4)
del numerical_columns[:]
del categorical_columns[:]

print("Programm Start: ", timestart)
timeend = datetime.datetime.now()
date_time = timeend.strftime("%d.%m.%Y %H:%M:%S")
print("Programm Finish:",date_time)
timedelta = round((timeend-timestart).total_seconds(), 2) 

r=(timeend-timestart) 
t=int(timedelta/60)
if timedelta-t*60 < 10:
    t2=":0" + str(int(timedelta-t*60))
else:
    t2=":" + str(int(timedelta-t*60))
txt="Programm Dauer: 00:" + str(t) + t2 
print(txt)
print("Programm AsNiPa203.ipynb Ende")

#heading = Doc.add_heading('', level=3)
#run = heading.add_run("my heading")
#run.bold = True
#run.font.name = 'Times New Roman'
#run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
#run.font.size = Pt(16)
